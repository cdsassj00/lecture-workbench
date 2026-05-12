#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
GIST 학위논문 양식 기반 학술논문 자동 생성기.

주제: 과학기술은 어떻게 경제적 가치가 되는가
     ― 이론, 메커니즘, 측정 및 한국에의 함의 ―

양식:
  - A4, 여백 상하 30mm / 좌우 25mm
  - 본문 한글 함초롬바탕(없으면 맑은 고딕/굴림), 영문 Times New Roman
  - 본문 11pt, 줄간격 2.0 (GIST 학위논문 규정)
  - 장(Chapter) 제목 18pt, 절(Section) 14pt, 항(Subsection) 12pt
  - 페이지 번호: 본문 아라비아 숫자, 서두 로마자
  - 표지 → 인준서 → 국문초록 → English Abstract → 목차 → 표/그림 목차
    → 본문 7개 장 → 참고문헌 → 부록 → 감사의 글 → 약력
"""

from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# ───────────────────────── 폰트/스타일 설정 ─────────────────────────
KOR_FONT = "맑은 고딕"
ENG_FONT = "Times New Roman"
BODY_PT = 11
TITLE_PT = 18
SECTION_PT = 14
SUBSECTION_PT = 12
LINE_SPACING = 2.0


def set_run_font(run, size_pt=BODY_PT, bold=False, italic=False, color=None):
    """런(Run)의 한·영 폰트와 속성을 동시에 지정한다."""
    run.font.name = ENG_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), KOR_FONT)
    rFonts.set(qn("w:ascii"), ENG_FONT)
    rFonts.set(qn("w:hAnsi"), ENG_FONT)


def add_paragraph(doc, text, *, size=BODY_PT, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_first=True,
                  spacing=LINE_SPACING, space_after=4, space_before=0,
                  color=None):
    """일반 단락 추가."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_first:
        pf.first_line_indent = Cm(1.0)
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1, *, page_break=False):
    """장/절/항 제목 추가. level 1=장, 2=절, 3=항."""
    if page_break and level == 1:
        doc.add_page_break()
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        size = TITLE_PT
        pf.space_before = Pt(24)
        pf.space_after = Pt(18)
        align = WD_ALIGN_PARAGRAPH.LEFT
        bold = True
    elif level == 2:
        size = SECTION_PT
        pf.space_before = Pt(14)
        pf.space_after = Pt(10)
        align = WD_ALIGN_PARAGRAPH.LEFT
        bold = True
    else:
        size = SUBSECTION_PT
        pf.space_before = Pt(8)
        pf.space_after = Pt(6)
        align = WD_ALIGN_PARAGRAPH.LEFT
        bold = True
    pf.line_spacing = 1.5
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return p


def add_quote(doc, text, source=""):
    """인용문 단락."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.right_indent = Cm(1.0)
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("“" + text + "”")
    set_run_font(run, size_pt=BODY_PT, italic=True)
    if source:
        run2 = p.add_run(f"  — {source}")
        set_run_font(run2, size_pt=BODY_PT - 1)
    return p


def add_table_caption(doc, num, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"<표 {num}> {caption}")
    set_run_font(run, size_pt=BODY_PT, bold=True)


def add_figure_caption(doc, num, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[그림 {num}] {caption}")
    set_run_font(run, size_pt=BODY_PT, bold=True)


def add_table(doc, header, rows, col_widths_cm=None):
    """간단한 3선표(상단·헤더하단·하단 굵은 선) 스타일 테이블."""
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = False
    # 헤더
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(h)
        set_run_font(run, size_pt=BODY_PT - 1, bold=True)
    # 데이터
    for r, row in enumerate(rows, start=1):
        cells = t.rows[r].cells
        for c, v in enumerate(row):
            cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[c].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(str(v))
            set_run_font(run, size_pt=BODY_PT - 1)
    # 컬럼 폭
    if col_widths_cm:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths_cm[i])
    # 3선표: 상단/헤더하단/하단만 굵은 선
    tbl = t._tbl
    tblBorders = OxmlElement("w:tblBorders")
    for edge, sz in (("top", 12), ("bottom", 12), ("insideH", 4)):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    for edge in ("left", "right", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr = tbl.tblPr
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return t


def set_page(section):
    """A4, 여백 GIST 규정."""
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(30)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)


def add_page_number(section, fmt="decimal", start=1):
    """페이지 번호를 푸터에 추가. fmt: 'decimal' 또는 'lowerRoman'."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    pgNumType.set(qn("w:start"), str(start))

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 기존 콘텐츠 제거
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    set_run_font(run, size_pt=BODY_PT)


# ───────────────────────── 본문 시작 ─────────────────────────
doc = Document()

# 전역 스타일
normal = doc.styles["Normal"]
normal.font.name = ENG_FONT
normal.font.size = Pt(BODY_PT)
rPr = normal.element.get_or_add_rPr()
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:eastAsia"), KOR_FONT)
rFonts.set(qn("w:ascii"), ENG_FONT)
rFonts.set(qn("w:hAnsi"), ENG_FONT)
rPr.append(rFonts)

# 섹션 1: 표지 (페이지 번호 없음)
set_page(doc.sections[0])

# ───── 표지 ─────
for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.15
run = p.add_run("박사학위논문")
set_run_font(run, size_pt=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ph.D. Dissertation")
set_run_font(run, size_pt=12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run("과학기술은 어떻게 경제적 가치가 되는가")
set_run_font(run, size_pt=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.3
run = p.add_run("― 이론, 메커니즘, 측정, 그리고 한국에의 함의 ―")
set_run_font(run, size_pt=14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.3
run = p.add_run("How Science and Technology Become Economic Value:")
set_run_font(run, size_pt=14, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.3
run = p.add_run("Theory, Mechanisms, Measurement, and Implications for Korea")
set_run_font(run, size_pt=14, italic=True)

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run("저   자   명")
set_run_font(run, size_pt=14, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("기술경영학과")
set_run_font(run, size_pt=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("광주과학기술원")
set_run_font(run, size_pt=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Gwangju Institute of Science and Technology")
set_run_font(run, size_pt=11, italic=True)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"{datetime.now().year}")
set_run_font(run, size_pt=14, bold=True)

# ───── 인준서 ─────
doc.add_page_break()
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("본 논문을 박사학위논문으로 인준함")
set_run_font(run, size_pt=14, bold=True)
for _ in range(3):
    doc.add_paragraph()
for label in [
    "심사위원장              (인)",
    "심사위원                (인)",
    "심사위원                (인)",
    "심사위원                (인)",
    "심사위원                (인)",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(label)
    set_run_font(run, size_pt=13)

# 섹션 2: 서두 (로마자 페이지 번호)
front_section = doc.add_section(WD_SECTION.NEW_PAGE)
set_page(front_section)
add_page_number(front_section, fmt="lowerRoman", start=1)

# ───── 국문초록 ─────
add_heading(doc, "국문초록", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("과학기술은 어떻게 경제적 가치가 되는가")
set_run_font(run, size_pt=14, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("― 이론, 메커니즘, 측정, 그리고 한국에의 함의 ―")
set_run_font(run, size_pt=12)
doc.add_paragraph()

abstract_ko = [
    "본 논문은 과학기술이 어떠한 경로와 메커니즘을 통해 경제적 가치(economic value)로 전환되는가라는 근본적 질문에 대한 학제적 답을 시도한다. 산업혁명 이래 인류의 경제성장은 자본의 축적이나 노동의 양적 증가보다 기술진보(technical progress)에 의해 더 크게 좌우되어 왔다는 것이 다수의 실증연구로 확인되어 왔지만, 정작 ‘과학적 지식이 어떻게 시장에서 가격을 갖는 재화·서비스로 구체화되는가’에 대한 통합적 설명은 여전히 충분하지 않다.",
    "이에 본 연구는 첫째, 슘페터(Schumpeter)의 혁신·창조적 파괴(creative destruction)와 솔로우(Solow) 잔차로 대표되는 신고전파 성장이론, 그리고 로머(Romer, 1990)·아기온–호윗(Aghion–Howitt, 1992)의 내생적 성장이론에 이르기까지의 이론적 계보를 정리한다. 둘째, 지식이 비경합성(non-rivalry)과 부분적 배제성(partial excludability)을 가진 재화임을 명확히 하고, 이러한 속성이 시장실패와 정부개입의 정당성, 그리고 지식의 사회적 수익률이 사적 수익률을 상회하는 현상을 설명함을 보인다. 셋째, 과학기술 → 발명 → 혁신 → 사업화 → 시장가치라는 가치사슬을 단계별로 분해하고, 각 단계에서 작용하는 제도적·조직적 요인—흡수역량(absorptive capacity), 기술이전, 트리플 헬릭스(triple helix), 국가혁신체제, 기업가형 국가—을 검토한다.",
    "측정 영역에서는 OECD Frascati Manual과 Oslo Manual을 토대로 한 R&D 통계, 특허·인용 분석, GERD(Gross Domestic Expenditure on R&D), R&D 자본화 이후의 GDP 계산법, 그리고 무형자산 회계 등 다층적 지표체계를 비교한다. 사례연구로는 1980년 미국의 Bayh-Dole법이 촉발한 대학 기술이전의 폭발적 증가(1996–2020년 누적 18,000여 개의 스핀오프 창업, 1조 9천억 달러의 산업산출 기여), 인터넷·GPS·터치스크린·mRNA 백신 등 정부 R&D가 후행적 민간 가치창출의 토대가 된 사례, 그리고 한국의 ICT·반도체·이차전지·AI 산업의 발전 경로를 분석한다.",
    "마지막으로 한국적 함의를 논한다. 한국은 GDP 대비 R&D 비중에서 세계 최상위권을 유지해 왔으나, 정부 R&D 예산의 급격한 변동, ‘추격형(catch-up)’에서 ‘선도형(first-mover)’으로의 전환 지연, 기초연구·응용연구·개발연구 간 균형, KISTEP의 평가체계, 그리고 임무지향형(mission-oriented) 혁신정책의 설계 등에서 구조적 과제를 안고 있다. 본 논문은 이러한 과제를 해소하기 위한 정책적 제언으로 (1) 연구개발 투자의 안정성과 다년도 예산제, (2) 대학·출연(연)·기업 간 인력과 지식의 양방향 유동성 강화, (3) 사회문제 해결형 임무지향 R&D의 거버넌스 정비, (4) 기술의 사회적 수용성과 형평성을 함께 고려하는 ‘책임 있는 혁신(responsible innovation)’의 제도화를 제시한다.",
    "본 연구의 학술적 기여는 다음과 같다. 첫째, 그동안 별개로 다루어진 경제성장 이론, 혁신연구, 기술이전, 측정 방법론을 하나의 통합적 프레임으로 결합하였다는 점이다. 둘째, 과학기술의 경제적 가치 전환을 단순한 ‘R&D → GDP’ 인과로 환원하지 않고 제도·인지·네트워크 요인을 명시적으로 통합하였다는 점이다. 셋째, 최근 인공지능과 디지털 전환이 가져온 가치창출 경로의 변화(데이터의 비경합성, 플랫폼의 양면시장 효과 등)를 기존 이론 체계 안에 자리매김하였다는 점이다.",
]
for para in abstract_ko:
    add_paragraph(doc, para)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("주제어: ")
set_run_font(run, size_pt=BODY_PT, bold=True)
run = p.add_run(
    "과학기술, 혁신, 경제적 가치, 내생적 성장이론, 창조적 파괴, R&D, "
    "기술이전, 트리플 헬릭스, 국가혁신체제, Bayh-Dole, KISTEP, "
    "임무지향형 혁신정책, 책임 있는 혁신"
)
set_run_font(run, size_pt=BODY_PT)

# ───── English Abstract ─────
add_heading(doc, "Abstract", level=1, page_break=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How Science and Technology Become Economic Value:")
set_run_font(run, size_pt=14, bold=True, italic=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Theory, Mechanisms, Measurement, and Implications for Korea")
set_run_font(run, size_pt=12, italic=True)
doc.add_paragraph()

abstract_en = [
    "This dissertation provides an integrated, interdisciplinary account of the question: through what pathways and mechanisms does scientific and technological knowledge become economic value? Although economists since Robert Solow have empirically established that technical progress, rather than capital deepening alone, drives long-run growth, a unified theory of how scientific ideas crystallize into priced goods and services in markets remains incomplete.",
    "We synthesize the lineage from Schumpeter's creative destruction and Solow's residual, through the new endogenous growth theory of Romer (1986, 1990) and Aghion and Howitt (1992), to the neo-Schumpeterian and evolutionary tradition of Nelson, Winter, Lundvall, and Freeman. We argue that the non-rivalry and partial excludability of knowledge are the fundamental properties that explain (i) increasing returns at the aggregate level, (ii) the gap between social and private rates of return, and (iii) the standing justification for public R&D investment.",
    "We then decompose the value chain from science to market—invention, innovation, diffusion, and commercialization—and examine the institutional and organizational mediators of each transition: absorptive capacity, technology transfer offices, the triple helix of university–industry–government relations, national innovation systems, and the entrepreneurial state. On measurement, we compare the OECD Frascati and Oslo Manuals, patent-citation analysis, gross domestic expenditure on R&D (GERD), R&D capitalization in national accounts, and intangible-asset accounting.",
    "Case studies span the United States' Bayh-Dole Act (1980)—which has produced more than 18,000 university spinoffs, 141,000 U.S. patents, and 1.9 trillion U.S. dollars of industrial output between 1996 and 2020—public-investment success stories such as the internet, GPS, touchscreen, and mRNA vaccines, and Korea's developmental trajectory in ICT, semiconductors, batteries, and artificial intelligence. We finally derive policy implications for Korea: multi-year R&D budgeting, bidirectional mobility between universities, government research institutes, and firms, mission-oriented innovation governance, and the institutionalization of responsible innovation.",
    "The dissertation makes three contributions. First, it bridges growth theory, innovation studies, technology transfer, and measurement methodology in a single framework. Second, it relocates the science-to-value transformation from a simple R&D-to-GDP causation to a multi-layered system involving institutions, cognition, and networks. Third, it integrates the recent shift in value creation associated with artificial intelligence and the digital transition—data non-rivalry, platform two-sidedness, and recombinant innovation—into the established theoretical edifice.",
]
for para in abstract_en:
    add_paragraph(doc, para)

doc.add_paragraph()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run("Keywords: ")
set_run_font(run, size_pt=BODY_PT, bold=True)
run = p.add_run(
    "science and technology, innovation, economic value, endogenous growth, "
    "creative destruction, R&D, technology transfer, triple helix, "
    "national innovation systems, Bayh-Dole, KISTEP, mission-oriented "
    "innovation policy, responsible innovation"
)
set_run_font(run, size_pt=BODY_PT, italic=True)

# ───── 목차 ─────
add_heading(doc, "목   차", level=1, page_break=True)

toc_entries = [
    ("국문초록", "i"),
    ("Abstract", "iii"),
    ("목차", "v"),
    ("표 목차", "viii"),
    ("그림 목차", "ix"),
    ("", ""),
    ("제1장  서론", "1"),
    ("  1.1  연구의 배경 및 필요성", "1"),
    ("  1.2  연구의 목적과 연구 문제", "4"),
    ("  1.3  연구의 범위와 방법", "6"),
    ("  1.4  논문의 구성", "8"),
    ("", ""),
    ("제2장  이론적 배경: 성장·혁신·지식에 대한 경제학적 사유", "10"),
    ("  2.1  고전·신고전 성장이론과 ‘기술’의 자리", "10"),
    ("  2.2  슘페터의 혁신 이론과 창조적 파괴", "14"),
    ("  2.3  솔로우 잔차와 ‘설명되지 않는 성장’", "17"),
    ("  2.4  로머의 내생적 성장이론: 지식의 비경합성", "19"),
    ("  2.5  아기온–호윗의 슘페터적 성장 모형", "23"),
    ("  2.6  진화경제학과 국가혁신체제(NIS)", "25"),
    ("  2.7  마추카토의 기업가형 국가론", "28"),
    ("  2.8  소결: 통합적 분석틀의 단초", "30"),
    ("", ""),
    ("제3장  과학기술이 경제적 가치로 전환되는 메커니즘", "32"),
    ("  3.1  지식–발명–혁신–확산–가치의 가치사슬", "32"),
    ("  3.2  기술이전과 사업화의 미시 메커니즘", "36"),
    ("  3.3  트리플 헬릭스: 대학–산업–정부의 상호작용", "40"),
    ("  3.4  흡수역량과 조직 학습", "43"),
    ("  3.5  네트워크 효과, 플랫폼, 양면시장", "46"),
    ("  3.6  재조합적 혁신과 일반목적기술(GPT)", "49"),
    ("", ""),
    ("제4장  경제적 가치의 측정", "52"),
    ("  4.1  R&D 통계의 국제 표준: Frascati Manual", "52"),
    ("  4.2  혁신활동의 측정: Oslo Manual", "55"),
    ("  4.3  특허·인용·논문계량의 지표학", "57"),
    ("  4.4  GDP 산정에서의 R&D 자본화", "60"),
    ("  4.5  무형자산 회계와 기업가치", "63"),
    ("  4.6  사회적 수익률 vs. 사적 수익률", "65"),
    ("", ""),
    ("제5장  사례 연구", "68"),
    ("  5.1  미국 Bayh-Dole 체제와 대학 기술이전", "68"),
    ("  5.2  공공 R&D가 잉태한 시장: 인터넷·GPS·터치스크린·mRNA", "73"),
    ("  5.3  한국 ICT 산업의 도약 경로", "77"),
    ("  5.4  반도체: 자본·기술·정책의 합주", "81"),
    ("  5.5  AI/LLM과 지능의 산업화", "85"),
    ("", ""),
    ("제6장  한국에의 정책적 함의", "89"),
    ("  6.1  한국 R&D 투자의 구조와 변동성", "89"),
    ("  6.2  KISTEP과 평가체계", "92"),
    ("  6.3  임무지향형 혁신정책의 한국적 적용", "95"),
    ("  6.4  글로벌 비교: 미·EU·일·중·이스라엘", "98"),
    ("  6.5  책임 있는 혁신과 사회적 수용성", "101"),
    ("", ""),
    ("제7장  결론", "104"),
    ("  7.1  연구의 요약", "104"),
    ("  7.2  학술적·정책적 시사점", "106"),
    ("  7.3  연구의 한계와 향후 과제", "108"),
    ("", ""),
    ("참고문헌", "110"),
    ("부록 A. 주요 개념·약어 정리", "118"),
    ("부록 B. 추가 통계 표", "120"),
    ("감사의 글", "122"),
    ("약   력", "123"),
]
for label, page in toc_entries:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.6
    pf.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)  # leader=dots
    if label:
        run = p.add_run(label + "\t" + page)
        set_run_font(run, size_pt=BODY_PT)
    else:
        run = p.add_run("")
        set_run_font(run, size_pt=BODY_PT)

# ───── 표 목차 ─────
add_heading(doc, "표 목차", level=1, page_break=True)
tables_index = [
    ("표 2.1", "주요 경제성장 이론과 ‘기술’의 위치", "16"),
    ("표 2.2", "지식의 경제적 속성: 사적재·공공재·클럽재 비교", "22"),
    ("표 3.1", "기술이전 채널과 그 특성", "39"),
    ("표 3.2", "흡수역량의 구성 요소(Cohen & Levinthal, 1990)", "45"),
    ("표 4.1", "Frascati Manual의 R&D 분류(2015 개정)", "54"),
    ("표 4.2", "Oslo Manual에 의한 혁신의 4유형", "56"),
    ("표 4.3", "주요국 GERD/GDP 비중(2022)", "62"),
    ("표 4.4", "R&D 사회적 수익률 추정치 메타분석", "67"),
    ("표 5.1", "Bayh-Dole 이후 미국 대학 기술이전 성과(1996–2020)", "71"),
    ("표 5.2", "한국 ICT 5대 품목의 세계시장 점유율 변화", "80"),
    ("표 6.1", "한국 정부 R&D 예산 추이(2018–2025)", "91"),
    ("표 6.2", "주요국 임무지향형 혁신정책 비교", "97"),
]
for n, c, pg in tables_index:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    run = p.add_run(f"{n}  {c}\t{pg}")
    set_run_font(run, size_pt=BODY_PT)

# ───── 그림 목차 ─────
add_heading(doc, "그림 목차", level=1, page_break=True)
figs_index = [
    ("그림 2.1", "솔로우 모형에서의 정상상태와 기술진보", "18"),
    ("그림 2.2", "Romer(1990) 모형의 3부문 구조", "21"),
    ("그림 2.3", "Aghion–Howitt 모형의 품질사다리", "24"),
    ("그림 3.1", "과학–기술–혁신–가치의 가치사슬", "33"),
    ("그림 3.2", "트리플 헬릭스 모형의 진화", "42"),
    ("그림 3.3", "플랫폼 양면시장 구조", "48"),
    ("그림 4.1", "GERD의 자금 흐름 (수행자–재원자 매트릭스)", "53"),
    ("그림 4.2", "한국 R&D 강도(GDP 대비) 추이", "61"),
    ("그림 5.1", "미국 대학 기술이전의 누적 성과", "72"),
    ("그림 5.2", "한국 ICT·반도체 수출 비중 변화", "82"),
    ("그림 6.1", "임무지향 R&D 거버넌스 개념도", "96"),
]
for n, c, pg in figs_index:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    run = p.add_run(f"{n}  {c}\t{pg}")
    set_run_font(run, size_pt=BODY_PT)

# ───────────────────────── 본문 섹션 ─────────────────────────
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
set_page(body_section)
add_page_number(body_section, fmt="decimal", start=1)

# ===================== 제1장 서론 =====================
add_heading(doc, "제1장  서   론", level=1)

add_heading(doc, "1.1  연구의 배경 및 필요성", level=2)

ch1_1 = [
    "21세기 세계경제에서 ‘과학기술’이라는 단어가 갖는 무게는 그 어느 시기보다 무겁다. 미국과 중국의 전략적 경쟁의 핵심에는 인공지능·반도체·바이오·우주 등 첨단 과학기술의 주도권 다툼이 자리하고, 유럽연합의 산업정책은 ‘디지털·녹색 전환(twin transition)’이라는 두 개의 기술적 변혁을 중심축으로 재편되고 있다. 국제통화기금(IMF), 세계은행, 경제협력개발기구(OECD) 등 주요 국제기구가 발간하는 연차 보고서들은 거의 예외 없이 기술혁신을 잠재성장률 제고의 가장 강력한 동인으로 지목한다. 우리나라 역시 1960년대 이후의 산업화, 1990년대 이후의 정보화, 2010년대 이후의 디지털 전환을 거치며 ‘과학기술 입국(立國)’을 국가 운영의 기본 전략으로 견지해 왔다.",
    "그럼에도 불구하고, ‘과학기술이 어떻게 경제적 가치가 되는가’라는 질문에 대한 우리의 답변은 종종 모호하다. 흔히 인용되는 ‘R&D 1단위 증가 시 GDP가 ○○ 증가한다’류의 회귀계수는 기술적 진보가 산업과 시장에서 가치로 구현되는 방대한 중간 과정을 검은 상자(black box)로 남겨 둔다. 솔로우(Solow, 1957)가 제시한 잔차(residual)는 자본·노동으로 설명되지 않는 산출 성장의 50% 이상을 ‘기술진보’라는 이름으로 정의하였으나, 그 ‘기술’이 무엇이며 어떻게 작동하는지는 모형 안에서 외생적인 채로 남겨졌다. 이후 로머(Romer, 1986; 1990)가 지식을 모형 내에서 내생화하면서 ‘잔차’가 ‘이론적 객체’가 되었지만, 그 거시적 명료성에 비해 미시·중간 단계의 메커니즘은 여전히 다른 분과(혁신연구, 경영학, 사회학)의 영역으로 남았다.",
    "한편 정책 현장에서는 더 절실한 질문들이 제기된다. 기초연구 1억 원을 늘리는 것과 응용연구 1억 원을 늘리는 것의 사회적 수익률은 어떻게 다른가? 대학에서 산출된 논문이 실제 산업의 생산성으로 이어지기까지 평균 몇 년이 걸리며, 그 경로에서 가장 큰 병목은 어디인가? 임무지향형 혁신정책(mission-oriented innovation policy)을 통해 사회문제 해결과 경제적 가치 창출을 동시에 달성하는 것은 가능한가? 인공지능과 같이 일반목적기술(general-purpose technology, GPT)의 성격을 갖는 신기술의 등장 앞에서, 국가는 어떤 종류의 투자를 어떤 거버넌스로 수행해야 하는가? 이러한 질문에 답하기 위해서는 거시 성장이론, 미시 혁신이론, 측정 방법론, 그리고 한국적 제도 분석이 통합적으로 동원되어야 한다.",
    "본 논문의 출발점은 다음과 같은 인식이다. 첫째, 과학기술의 경제적 가치 전환은 단순한 인과의 사슬이 아니라 다층적·다행위자적·진화적 시스템이다. 둘째, 그 시스템을 이해하려면 슘페터에서 로머에 이르는 경제학의 이론적 자원, 프리먼(Freeman)·룬드발(Lundvall)·넬슨(Nelson)으로 대표되는 혁신연구의 시스템적 접근, 그리고 OECD·UNESCO·세계은행이 정련해 온 측정 방법론을 결합해야 한다. 셋째, 한국은 ‘추격형 혁신(catch-up innovation)’의 정점에서 ‘선도형 혁신(first-mover innovation)’으로의 전환이라는 역사적 과제 앞에 서 있으며, 이 과제에 대한 답은 이론적 통합과 정책적 구체성의 결합 없이는 불가능하다.",
    "본 연구는 따라서 학술적 야심과 정책적 절박함 사이의 다리를 놓고자 한다. 단순한 문헌 종합이나 모형 검증을 넘어서, 과학기술이 경제적 가치로 전환되는 ‘전 과정’을 단일한 분석틀로 묶고, 한국이 직면한 구체적인 정책 선택지에 대해 이론적 근거가 분명한 제언을 제시하는 것이 본 논문의 궁극적 목적이다.",
]
for p in ch1_1:
    add_paragraph(doc, p)

add_heading(doc, "1.2  연구의 목적과 연구 문제", level=2)
ch1_2 = [
    "본 논문의 핵심 목적은 ‘과학기술이 경제적 가치로 전환되는 메커니즘’에 대한 통합적·다층적 분석틀을 제시하고, 그 분석틀을 한국의 혁신체제에 적용하여 정책적 함의를 도출하는 데 있다. 이를 위해 본 연구는 다음의 네 가지 연구 문제(research questions)를 설정한다.",
    "연구 문제 1. 과학기술과 경제성장의 관계를 설명하는 경제학의 이론적 계보는 어떻게 발전해 왔으며, 각 이론은 ‘기술’을 어떤 위상으로 다루고 있는가? 본 질문은 고전 경제학에서부터 솔로우 모형, 로머의 R&D 기반 내생적 성장, 아기온–호윗의 슘페터적 성장, 그리고 진화경제학적 NIS 접근에 이르기까지의 이론적 흐름을 비판적으로 조망한다.",
    "연구 문제 2. 거시 차원의 ‘기술–성장’ 관계는 미시 차원에서 어떠한 구체적 메커니즘으로 매개되는가? 본 질문은 지식–발명–혁신–확산의 가치사슬, 기술이전 채널, 트리플 헬릭스, 흡수역량, 네트워크 효과 등 중간 메커니즘을 다룬다.",
    "연구 문제 3. 과학기술의 경제적 가치는 무엇으로 어떻게 측정되며, 각 측정 방식의 한계와 보완책은 무엇인가? 본 질문은 Frascati·Oslo Manual, 특허·인용 분석, R&D 자본화, 무형자산 회계 등을 다층적으로 비교한다.",
    "연구 문제 4. 한국의 과학기술–경제 가치 시스템은 어떤 구조적 특징과 약점을 가지며, 어떤 방향의 정책 재설계가 요구되는가? 본 질문은 KISTEP의 평가체계, 정부 R&D 예산 거버넌스, 임무지향 혁신정책의 적용, 책임 있는 혁신 등을 다룬다.",
    "이 네 가지 질문은 일종의 ‘이론 → 메커니즘 → 측정 → 정책’의 논리적 사다리를 형성한다. 본 논문의 각 장은 이 사다리의 한 단(段)에 대응하며, 전체적으로 ‘과학기술이 경제적 가치가 되는 길’을 입체적으로 조명할 것이다.",
]
for p in ch1_2:
    add_paragraph(doc, p)

add_heading(doc, "1.3  연구의 범위와 방법", level=2)
ch1_3 = [
    "본 연구의 분석 범위는 다음과 같이 설정한다. 시간적으로는 산업혁명(18세기 말)에서 출발하여 2025년 현재까지를 포괄하되, 분석의 무게중심은 제2차 세계대전 이후의 현대 경제·정책 체제에 둔다. 공간적으로는 OECD 회원국을 중심으로 한 글로벌 비교를 수행하되, 한국에 대해서는 미시 정책 단위까지 구체적으로 다룬다. 산업적으로는 ICT·반도체·바이오·이차전지·인공지능을 핵심 사례 산업으로 선택하였다.",
    "방법론적으로 본 연구는 문헌 종합(literature synthesis), 이론 분석(theoretical analysis), 사례 비교 분석(comparative case study), 그리고 통계 자료 검토(secondary data review)를 결합한다. 1차적으로는 NBER Working Paper, Journal of Political Economy, American Economic Review, Research Policy, Strategic Management Journal 등 학술지의 논문, 그리고 OECD·UNESCO·세계은행·IMF·EU 집행위원회·미국 NSF·한국 KISTEP·과학기술정책연구원(STEPI) 등의 공식 보고서를 활용한다. 2차적으로는 통계 데이터로 OECD Main Science and Technology Indicators(MSTI), UNESCO UIS R&D Database, 한국 국가과학기술지식정보서비스(NTIS)의 공개 자료를 인용한다.",
    "본 연구는 새로운 계량경제학적 회귀분석을 수행하기보다는, 기존의 다양한 실증 결과를 이론적 프레임 안에 재배치하고 해석하는 ‘이론적 통합 연구(theoretical integration study)’의 성격을 갖는다. 이는 학위논문의 통상적 형식인 ‘실증분석 1편’과 다른 선택이지만, 본 연구의 주제—‘메커니즘의 종합’—를 위해 적절하다고 판단된다. 다만 각 주장의 신뢰성을 확보하기 위해, 인용된 모든 실증 추정치는 메타분석 또는 다수 연구의 수렴값을 우선한다.",
]
for p in ch1_3:
    add_paragraph(doc, p)

add_heading(doc, "1.4  논문의 구성", level=2)
ch1_4 = [
    "본 논문은 총 7개 장으로 구성된다. 제1장은 본 서론이다. 제2장은 과학기술과 경제성장에 관한 경제학의 이론적 계보를 정리한다. 제3장은 과학기술이 경제적 가치로 전환되는 미시·중간 메커니즘을 다룬다. 제4장은 그 가치를 측정하는 다양한 방법론을 비교한다. 제5장은 미국 Bayh-Dole 체제, 공공 R&D의 산업화 사례, 한국 ICT·반도체·AI 산업의 발전 경로를 사례 연구한다. 제6장은 한국의 R&D 거버넌스와 임무지향 혁신정책을 분석하고 구체적 정책 제언을 도출한다. 제7장은 본 연구의 요약과 한계, 그리고 향후 연구 과제를 제시한다. 부록에는 주요 약어·개념의 정리와 보충 통계표를 포함한다.",
]
for p in ch1_4:
    add_paragraph(doc, p)

# ===================== 제2장 =====================
add_heading(doc, "제2장  이론적 배경: 성장·혁신·지식에 대한 경제학적 사유",
            level=1, page_break=True)

add_heading(doc, "2.1  고전·신고전 성장이론과 ‘기술’의 자리", level=2)
ch2_1 = [
    "근대 경제학의 출발선에 선 아담 스미스(Adam Smith, 1776)는 『국부론』의 서두에서 분업과 기계의 발명이 노동생산성을 향상시키는 가장 큰 원인임을 명시하였다. 그가 든 핀 공장의 사례는 단순한 기술적 진보가 아니라 ‘분업의 심화 → 기계화 → 시장 확대 → 다시 분업의 심화’라는 누적적 인과 구조를 함의한다. 이러한 통찰은 19세기 마르크스(Marx)의 ‘기계와 대공업’ 논의로 이어졌고, 마르크스는 자본주의가 끊임없이 생산력을 혁명함으로써만 존속한다고 보았다. 그러나 신고전파 경제학이 자리 잡으면서, 이러한 동태적 통찰은 균형 분석의 정태성에 가려지는 경향을 보였다.",
    "신고전파의 첫 본격적 성장 모형인 솔로우(Solow, 1956)와 스완(Swan, 1956)의 모형은 자본축적과 노동성장률이 외생적으로 주어진 조건에서 일인당 산출의 동학을 분석하였다. 이 모형의 핵심 결론은 ‘자본축적만으로는 지속적 성장이 불가능하며, 장기 일인당 성장률은 외생적 기술진보율과 같다’는 것이다. 즉, 솔로우 모형 안에서 ‘기술’은 인과적으로 가장 중요한 변수임에도 불구하고 그 결정 메커니즘이 분석되지 않는, 모형 외부의 ‘하늘에서 떨어지는 만나(manna from heaven)’로 남았다.",
    "이 외생성 가정은 1960~70년대 다양한 형태의 기술 매개변수화 시도—Arrow(1962)의 학습에 의한 행위(learning by doing), Uzawa(1965)의 인적자본 축적, Phelps(1966)의 R&D 변수화—를 거치면서 점차 약화되었으나, 본격적인 내생화는 1980년대 후반 로머의 ‘지식의 비경합성’ 정식화를 기다려야 했다. 본 절에서는 우선 솔로우 모형의 정식과 그 한계를 짚는다.",
    "솔로우 모형은 콥–더글러스 형태의 총생산함수 Y = K^α(AL)^(1−α)에서 출발한다. 여기에서 K는 자본, L은 노동, A는 ‘기술 수준’ 또는 ‘노동의 효율성(labor-augmenting technology)’으로 해석된다. 이 모형에서 일인당 산출 y = Y/(AL)은 자본–효율노동비 k = K/(AL)의 함수가 되고, k는 정상상태(steady state) k*에 수렴한다. 정상상태에서 일인당 산출 Y/L의 성장률은 g_A, 즉 기술진보율과 같다.",
    "이 결과의 정책적 함의는 다음과 같다. 첫째, 일인당 소득의 ‘수준’은 저축률·인구성장률·감가상각률 등에 의존하지만, 장기 ‘성장률’은 기술진보율 g_A에만 의존한다. 둘째, 따라서 한 나라의 장기 번영을 결정하는 가장 본질적 변수는 ‘기술’이며, 정책의 궁극적 표적은 g_A를 높이는 것이어야 한다. 셋째, 그러나 모형 안에서 g_A가 어떻게 결정되는지는 침묵하므로, 정책 처방의 구체성은 모형 바깥에서 보충되어야 한다. 솔로우 모형의 이러한 ‘불완전성’이 바로 이후 내생적 성장이론의 진입점이 되었다.",
    "한편 솔로우(1957)의 두 번째 논문은 미국 1909–1949년의 비농업 부문 데이터에 모형을 적용하여 ‘일인당 산출 성장의 약 87.5%가 기술진보로 설명된다’는 충격적인 추정 결과를 제시하였다. 이른바 ‘솔로우 잔차(Solow residual)’의 등장이다. 이후 다양한 계량경제학적 정련 작업에도 불구하고, 잔차가 차지하는 비중은 OECD 전체 평균으로 50% 이상, 다수의 신흥국에서도 30~50% 수준을 유지하였다(Easterly & Levine, 2001). ‘무엇이 이 잔차를 채우는가’가 이후 반세기 성장경제학의 핵심 의제가 되었다.",
]
for p in ch2_1:
    add_paragraph(doc, p)

add_heading(doc, "2.2  슘페터의 혁신 이론과 창조적 파괴", level=2)
ch2_2 = [
    "신고전파의 균형 분석과 대척점에서, 조지프 슘페터(Joseph A. Schumpeter)는 자본주의의 본질을 끊임없는 동태적 변화로 파악하였다. 그의 초기 저작 『경제 발전의 이론』(Theorie der wirtschaftlichen Entwicklung, 1912)은 정태적 일반균형과는 다른 ‘발전(Entwicklung)’ 개념을 도입하였다. 슘페터에게 발전이란 단순한 양적 성장이 아니라 새로운 결합(neue Kombinationen)을 통해 기존 질서를 깨뜨리는 비연속적 도약이며, 그 주체는 기업가(Unternehmer)이다.",
    "슘페터는 새로운 결합을 다섯 가지로 분류하였다. (1) 새로운 재화의 도입, (2) 새로운 생산방법의 도입, (3) 새로운 시장의 개척, (4) 새로운 원료·반제품 공급원의 발견, (5) 새로운 산업조직의 형성이 그것이다. 흥미롭게도 이 분류는 OECD의 Oslo Manual(2018)에서 정의하는 ‘제품 혁신, 공정 혁신, 마케팅 혁신, 조직 혁신’의 원형이다. 백 년 전의 슘페터적 분류가 오늘날 R&D·혁신 통계의 표준 분류 체계 안에 살아 있다는 사실은, 그의 이론적 통찰의 깊이를 방증한다.",
    "후기 저작 『자본주의, 사회주의, 민주주의』(1942)에서 슘페터는 ‘창조적 파괴(creative destruction)’라는 용어를 도입하였다. 그에 따르면, “자본주의 엔진을 작동시키고 유지하는 근본적 충격은, 자본주의 기업이 만들어내는 새로운 소비재, 새로운 생산 또는 운송 방법, 새로운 시장, 새로운 형태의 산업조직에서 온다.” 이 동력은 기존의 산업 구조를 ‘내부로부터’ 끊임없이 혁명함으로써, 낡은 것을 파괴하고 새로운 것을 창조한다. 이러한 비연속적 변화의 과정이야말로 자본주의의 본질이며, 정태적 가격 균형의 분석으로는 포착할 수 없다.",
    "슘페터의 통찰은 두 가지 점에서 본 연구의 출발점과 직결된다. 첫째, ‘과학기술이 어떻게 경제적 가치가 되는가’의 답은 가격 균형이 아니라 ‘새로운 결합’의 과정에서 찾아야 한다는 인식론적 입장이다. 둘째, 혁신은 본질적으로 파괴를 수반하므로, 그 사회적 비용과 편익의 비대칭성, 그리고 ‘누가 보상받고 누가 손실을 입는가’라는 정치경제학적 질문이 회피될 수 없다는 점이다. 이는 본 논문 제6장의 ‘책임 있는 혁신’ 논의와 직접 연결된다.",
    "슘페터의 영향은 1970년대 이후 ‘신슘페터 학파(neo-Schumpeterian school)’로 부활하였다. 넬슨과 윈터(Nelson & Winter, 1982)는 『경제변화의 진화론적 이론』에서 기업을 ‘루틴(routine)의 묶음’으로 보고, 기업 간 변이·선택·보존이라는 다윈적 메커니즘이 산업 진화를 추동한다고 주장하였다. 이 진화경제학적 전통은 이후 룬드발(Lundvall, 1992)과 프리먼(Freeman, 1987)의 국가혁신체제(National Innovation System, NIS) 개념으로 발전한다.",
    "또한 슘페터적 동학을 신고전 성장모형의 수학적 언어로 다시 쓰는 작업은 아기온과 호윗(Aghion & Howitt, 1992)의 ‘품질사다리(quality ladder)’ 모형에서 절정에 달한다. 이들의 모형에서는 R&D를 통해 기존 제품을 더 높은 품질로 대체하는 혁신가가 일정 기간 독점지대를 누리지만, 이는 곧 다음 혁신가에 의해 파괴된다. 이러한 시계열적 ‘파괴와 창조’의 반복이 거시적 성장률을 결정한다. 아기온–호윗 모형은 본 논문 2.5절에서 보다 자세히 다룬다.",
]
for p in ch2_2:
    add_paragraph(doc, p)

# 표 2.1
add_table_caption(doc, "2.1", "주요 경제성장 이론과 ‘기술’의 위치")
add_table(
    doc,
    header=["이론/모형", "대표 학자", "기술의 위치", "정책 함의"],
    rows=[
        ["고전 경제학", "Smith, Marx", "분업·기계의 누적적 발전",
         "시장 확대와 자본 축적"],
        ["솔로우 모형", "Solow(1956,1957)", "외생적 진보 g_A",
         "기술진보율 자체는 정책 외 변수"],
        ["AK 모형", "Romer(1986)",
         "자본의 사회적 한계생산 일정", "투자의 사회수익률에 주목"],
        ["내생적 성장 R&D 모형", "Romer(1990)",
         "R&D 부문 노동투입으로 결정", "R&D 보조·인적자본 정책"],
        ["슘페터 성장", "Aghion–Howitt(1992)",
         "품질사다리 위의 ‘파괴’와 ‘창조’", "경쟁정책·진입장벽 완화"],
        ["진화경제학", "Nelson–Winter(1982)",
         "루틴의 변이·선택·보존", "기업 다양성·실험의 보호"],
        ["국가혁신체제", "Lundvall(1992), Freeman(1987)",
         "제도·네트워크의 산물", "체제 차원 거버넌스"],
        ["기업가형 국가", "Mazzucato(2013)",
         "공공이 시장 형성을 주도", "미션 지향 공공투자"],
    ],
    col_widths_cm=[3.2, 3.5, 5.0, 4.0],
)

add_heading(doc, "2.3  솔로우 잔차와 ‘설명되지 않는 성장’", level=2)
ch2_3 = [
    "솔로우 잔차의 발견은 두 가지 상반된 반응을 낳았다. 한편에서는 ‘기술이 성장의 가장 큰 동인이다’라는 명료한 정책 메시지를 강화하였고, 다른 한편에서는 ‘우리가 측정 못 하는 모든 것을 기술이라 부른다’는 회의적 평가를 자극하였다. Griliches(1995)는 “잔차는 우리의 무지의 척도(measure of our ignorance)”라는 유명한 표현으로 후자의 입장을 요약하였다.",
    "잔차의 ‘내용’을 채우려는 시도는 크게 세 갈래로 진행되었다. 첫째, 인적자본의 명시적 도입이다. 만큐–로머–웨일(Mankiw, Romer & Weil, 1992)은 솔로우 모형에 인적자본 H를 추가하여 표본국 간 일인당 소득 격차의 약 80%를 설명할 수 있음을 보였다. 둘째, R&D 자본의 명시적 도입이다. Coe & Helpman(1995)은 국가별·산업별 R&D 스톡 변수가 총요소생산성(TFP)을 설명하는 강력한 변수임을 보이고, 무역을 통한 R&D 파급(R&D spillover) 효과를 정량화하였다. 셋째, 제도 변수의 도입이다. Acemoglu·Johnson·Robinson(2001) 등은 식민지 시기의 제도 형태가 현대의 일인당 소득 차이의 큰 부분을 설명한다고 주장하였다.",
    "이 세 갈래의 시도는 모두 솔로우 잔차의 ‘잉여(residual)’ 성격을 점차 줄여 왔지만, 동시에 새로운 문제를 제기하였다. 인적자본의 측정은 평균 교육연수에 의존하는데, 교육의 질과 분야 구성을 제대로 반영하지 못한다(Hanushek & Woessmann, 2008). R&D 스톡의 측정은 영구재고법(perpetual inventory method, PIM)에 의존하는데, 감가상각률의 선택이 결과를 좌우한다(Hall et al., 2010). 제도 변수의 외생성 문제는 도구변수 식별의 타당성 논쟁을 낳았다(Albouy, 2012). 즉, 잔차는 줄어들었으되, 측정 자체의 불확실성이 새로운 잔차로 자리바꿈한 셈이다.",
    "본 논문의 입장은 다음과 같다. 솔로우 잔차의 ‘크기’ 자체에 매달리기보다, 잔차가 작아진 만큼 그 ‘내용’을 다층적·동태적으로 해명하는 것이 더 생산적이다. 이는 거시 회귀의 한 자릿수 계수보다, 미시·중간 메커니즘의 풍부한 묘사를 우선하는 본 논문의 방법론적 선택과 일치한다.",
]
for p in ch2_3:
    add_paragraph(doc, p)

add_heading(doc, "2.4  로머의 내생적 성장이론: 지식의 비경합성", level=2)
ch2_4 = [
    "폴 로머(Paul M. Romer)의 1986년 논문 「Increasing Returns and Long-Run Growth」와 1990년 논문 「Endogenous Technological Change」는 현대 성장이론의 분기점이다. 그 핵심 통찰은 단순하지만 강력하다. 첫째, 지식은 비경합재(non-rival good)이다. 둘째, 지식의 생산에는 고정비용이 들지만 한 단위 더 생산에는 한계비용이 거의 0이다. 셋째, 이 두 속성이 결합하면 한 경제 전체로는 ‘규모에 대한 수익 증가(increasing returns to scale)’가 발생하며, 이것이 지속적 성장의 원천이 된다.",
    "로머(1990) 모형은 경제를 세 부문—최종재 부문, 중간재 부문, R&D 부문—으로 나눈다. 최종재는 노동, 인적자본의 일부, 그리고 다양한 중간재의 합으로 생산된다. 중간재는 R&D를 통해 새로 발명된 ‘설계도(blueprint)’를 기반으로 독점 생산된다. R&D 부문은 인적자본과 기존 지식 스톡에 비례하여 새로운 설계도를 산출한다. 이때 지식 스톡 A의 성장률은 R&D 부문에 투입된 인적자본 H_A에 비례한다. 그 결과, 정상상태에서의 일인당 산출 성장률 g_y = g_A는 H_A의 크기에 의해 결정된다.",
    "이 모형이 신고전 성장이론에 미친 충격은 두 가지였다. 첫째, ‘성장률’이 모형 안에서 결정된다. 즉, R&D 정책·인적자본 정책·지식재산권 보호 정책이 일인당 소득의 ‘수준’뿐 아니라 ‘성장률’에 영향을 미칠 수 있음이 이론적으로 정당화된 것이다. 둘째, 지식의 비경합성이 ‘시장실패’를 함의하므로, 사적 R&D 투자는 사회적 최적치에 미달한다. 따라서 R&D 보조금, 특허제도, 공공 연구투자 등 정부 개입이 정당화된다.",
    "Jones(1995)는 로머 모형의 ‘규모효과(scale effect)’—인구가 클수록 성장률이 높다는 함의—가 실증과 충돌함을 지적하고, 반(半)내생적 성장(semi-endogenous growth)을 제안하였다. 이 변형에서는 R&D의 양적 증가에 따른 수익이 체감하므로, 장기 성장률은 인구성장률에 비례하게 된다. 이후 ‘목표지향적 R&D’ 모형, ‘기술 도입과 이종 결합(recombination)’ 모형, ‘지식의 코드화 정도(codification)’ 모형 등 다양한 변형이 제안되었으나, 로머 1990 모형의 기본 구조는 여전히 분석의 기준점이다.",
    "지식의 비경합성은 본 논문의 핵심 메시지와 직결된다. 비경합성은 두 가지 정책적 귀결을 낳는다. 한편으로, 지식은 일단 생산되면 거의 영(零)에 가까운 비용으로 무한 복제·전파될 수 있으므로, 그 사회적 사용을 극대화하는 것이 효율적이다. 다른 한편으로, 발명에 든 고정비용을 회수할 사적 인센티브를 보호하지 않으면 처음부터 발명이 일어나지 않는다. 이 두 요구의 긴장(non-rivalry vs. excludability)이 지식재산권 제도, 공공 R&D, 오픈사이언스 정책의 본질적 트레이드오프이며, 본 논문 제4·6장의 측정·정책 논의의 사상적 토대가 된다.",
]
for p in ch2_4:
    add_paragraph(doc, p)

# 표 2.2
add_table_caption(doc, "2.2", "지식의 경제적 속성: 사적재·공공재·클럽재 비교")
add_table(
    doc,
    header=["속성", "사적재", "공공재", "클럽재", "지식(이상형)"],
    rows=[
        ["경합성", "있음", "없음", "없음(혼잡 한계 내)", "없음(거의)"],
        ["배제성", "있음", "없음", "있음(회비)", "부분적(특허·노하우)"],
        ["대표 예", "빵, 의류", "국방, 등대", "유료 도서관", "수식·알고리즘"],
        ["시장 공급", "효율적", "과소공급", "최적치 미달 가능", "사적 과소공급"],
        ["정책 처방", "최소 개입", "정부 직접 공급", "회원제·공동소유",
         "특허·R&D 보조·공공연구"],
    ],
    col_widths_cm=[3.0, 2.5, 2.5, 3.0, 4.5],
)

add_heading(doc, "2.5  아기온–호윗의 슘페터적 성장 모형", level=2)
ch2_5 = [
    "필립 아기온과 피터 호윗(Aghion & Howitt, 1992)은 슘페터의 창조적 파괴를 신고전 성장모형의 수학적 언어로 옮긴 결정적 작업을 수행하였다. 그들의 모형에서 각 산업은 ‘품질사다리(quality ladder)’의 한 단(段)에 있고, R&D에 성공한 혁신가는 한 단 위의 품질을 가진 중간재를 시장에 공급함으로써 일정 기간 독점지대를 누린다. 이 독점지대가 R&D의 보상이지만, 그 보상은 다음 혁신가의 등장과 함께 파괴된다. 이러한 ‘기대된 파괴’가 현재 혁신가의 기대수익을 감소시키지만(소위 ‘creative destruction effect’), 동시에 시장 전체로는 더 빈번한 혁신을 유도하여 평균 성장률을 높인다.",
    "이 모형의 정책 함의는 풍부하다. 첫째, 경쟁이 강할수록 혁신이 증가하는 ‘경쟁–혁신’의 U자형 관계가 도출된다(Aghion et al., 2005). 너무 약한 경쟁은 독점지대를 누리기 위한 진입을 약화하고, 너무 강한 경쟁은 혁신의 사후 보상을 침식한다. 둘째, 후발기업이 선도기업에 ‘추격하려는 동기(escape-competition effect)’가 클수록 혁신이 증가한다. 셋째, R&D 보조금의 효과는 산업의 기술 격차(technological gap)에 따라 다르다.",
    "아기온–호윗 모형은 본 논문 제5장의 사례 연구—특히 반도체·디스플레이·이차전지처럼 ‘선도와 추격’이 끊임없이 반복되는 산업—에 직접 적용된다. 또한 본 논문 제6장의 정책 논의에서, 공정거래법·시장지배적 지위 규제와 R&D 정책 간의 균형을 어떻게 잡을 것인가에 대한 이론적 기반을 제공한다.",
]
for p in ch2_5:
    add_paragraph(doc, p)

add_heading(doc, "2.6  진화경제학과 국가혁신체제(NIS)", level=2)
ch2_6 = [
    "1980년대 중후반, 신고전 성장이론과 다른 방향에서 새로운 흐름이 형성되었다. 크리스토퍼 프리먼(Christopher Freeman, 1987)은 일본의 경제적 약진을 분석하면서 ‘국가혁신체제(National Innovation System)’ 개념을 제시하였다. 같은 시기 룬드발(Lundvall, 1992)과 넬슨(Nelson, 1993) 역시 ‘체제로서의 혁신’이라는 관점을 공유하였다. NIS는 한 국가의 경제적 성과가 단순한 R&D 투입의 크기보다 R&D 수행 주체(기업·대학·정부연구소), 그들 간의 상호작용, 그리고 이를 지원하는 제도(금융·교육·법·문화)의 총체적 ‘시스템’에 의해 결정된다고 본다.",
    "NIS 관점의 가장 중요한 학술적 기여는 다음과 같다. 첫째, ‘동일한 R&D 강도’를 가진 두 국가가 매우 다른 혁신 성과를 보일 수 있음을 설명한다. 예컨대, 1990년대 일본과 미국은 GDP 대비 R&D 비중이 비슷했지만, 신생 산업의 부상 속도와 기업 진입·퇴출의 동학은 크게 달랐다. NIS는 이 차이를 ‘체제’의 차이로 설명한다. 둘째, 추격형 경제(catch-up economy)와 선도형 경제(frontier economy)에서 요구되는 혁신체제의 구성이 다르다는 통찰을 제공한다. 추격기에는 흡수·모방·역설계(reverse engineering)의 효율성이 중요하지만, 선도기에는 기초연구·실험적 시도·실패 허용의 제도적 기반이 더 중요하다.",
    "한국은 NIS 관점의 ‘대표 사례’로 자주 인용된다. 1960년대 후반의 한국과학기술연구원(KIST) 설립, 1970년대의 중화학공업 R&D 투자, 1980년대의 대기업 자체 R&D 도약, 1990년대의 정부출연(연) 확장과 대학원 중심 대학 육성, 2000년대의 BK21 사업과 신성장동력 정책은 모두 ‘체제로서의 혁신 정책’의 명료한 사례이다. 본 논문 제5·6장에서 한국 사례를 다룰 때 NIS의 어휘를 빈번하게 사용할 것이다.",
    "진화경제학의 어휘(루틴, 변이, 선택, 보존)와 NIS의 어휘(주체, 상호작용, 제도)는 상호 보완적이다. 전자는 미시 단위(기업·연구실)의 진화를 다루고, 후자는 거시 단위(국가·지역)의 시스템을 다룬다. 본 논문은 두 어휘를 모두 활용하면서, 그 사이의 ‘중간 단위(meso-level)’—산업 클러스터, 기술이전 네트워크, 대학–기업 연구 컨소시엄—를 명시적으로 다룬다.",
]
for p in ch2_6:
    add_paragraph(doc, p)

add_heading(doc, "2.7  마추카토의 기업가형 국가론", level=2)
ch2_7 = [
    "마리아나 마추카토(Mariana Mazzucato)의 『기업가형 국가(The Entrepreneurial State)』(2013)는 기존의 ‘민간 주도 혁신/공공 보조적 역할’ 통념을 정면으로 반박한 저작이다. 그녀의 주장은 다음과 같다. 첫째, 인터넷·GPS·터치스크린·음성인식·이산자율 자율주행 알고리즘 등 21세기 핵심 기술의 ‘위험 자본(risk capital)’은 대부분 공공 부문이 부담하였다. 둘째, 공공 부문은 단순한 시장실패 보정자가 아니라, ‘시장을 형성(market-shaping)’하고 ‘방향성(directionality)’을 제시하는 능동적 행위자였다. 셋째, 그럼에도 불구하고 그 결과의 사적 전유(privatization of returns)는 충분히 제어되지 않았고, 이로 인해 ‘공공이 위험을 부담하고 사적 부문이 보상을 가져가는’ 비대칭이 누적되었다.",
    "마추카토의 통찰은 본 논문에 두 가지 직접적 함의를 가진다. 첫째, ‘과학기술이 어떻게 경제적 가치가 되는가’의 답은 ‘민간기업의 R&D’와 ‘공공의 기초연구 보조’라는 이분법으로 환원되지 않는다. 공공 부문은 종종 ‘파괴적 혁신(disruptive innovation)’의 가장 큰 자본주이며, 그 역할의 정당한 평가와 보상 메커니즘이 정책 설계의 중심 과제가 되어야 한다. 둘째, 임무지향형 혁신정책(mission-oriented innovation policy)이라는 새로운 패러다임의 이론적 기반이 마련된다. 이는 본 논문 제6장의 핵심 정책 논의로 직결된다.",
    "비판도 적지 않다. Mingardi(2015) 등은 공공 부문의 ‘성공 사례’만 부각하는 선택 편의(selection bias)와, 공공 R&D의 막대한 실패 비용에 대한 침묵을 지적한다. 마추카토 자신도 후속 저작(Mission Economy, 2021; Public Purpose, 2024)에서 이러한 비판을 일부 수용하며, ‘공공 부문도 실패한다. 그러나 그 실패의 학습 가치를 어떻게 사회 전체로 환원할 것인가가 관건’이라는 보다 정제된 입장을 제시하였다.",
]
for p in ch2_7:
    add_paragraph(doc, p)

add_heading(doc, "2.8  소결: 통합적 분석틀의 단초", level=2)
ch2_8 = [
    "이상의 이론적 검토를 종합하면, ‘과학기술이 어떻게 경제적 가치가 되는가’에 대한 단일하고 환원적인 답은 존재하지 않는다. 솔로우 모형은 거시적 윤곽을, 로머 모형은 지식의 속성을, 아기온–호윗 모형은 경쟁과 진입의 동학을, 넬슨–윈터 진화경제학과 NIS는 제도와 시스템의 효과를, 마추카토의 기업가형 국가론은 공공의 능동적 역할을 각각 조명한다. 본 논문이 제3장 이후 다룰 ‘메커니즘–측정–정책’의 분석은 이 모든 이론적 자원의 ‘합성(synthesis)’에 기반한다.",
    "구체적으로, 본 논문은 다음의 통합적 분석틀을 채택한다. (1) 가치 창출의 거시적 동력은 R&D 투자·인적자본·제도의 결합이다. (2) 이 동력이 미시적으로 작동하는 채널은 발명·혁신·확산의 가치사슬이며, 각 채널에서 트리플 헬릭스의 행위자들이 상호작용한다. (3) 이 상호작용은 비경합성과 부분적 배제성이라는 지식의 속성에 의해 본질적으로 제약된다. (4) 따라서 측정과 정책은 이 다층적 시스템 전체를 시야에 두어야 하며, 어느 한 단계의 ‘부분적 최적화’가 전체 시스템의 손상을 낳을 수 있음을 항상 경계해야 한다.",
]
for p in ch2_8:
    add_paragraph(doc, p)

# ===================== 제3장 =====================
add_heading(doc, "제3장  과학기술이 경제적 가치로 전환되는 메커니즘",
            level=1, page_break=True)

add_heading(doc, "3.1  지식–발명–혁신–확산–가치의 가치사슬", level=2)
ch3_1 = [
    "과학적 지식이 경제적 가치로 전환되기까지의 경로는 단일한 직선이 아니라 복합적인 네트워크의 형태를 띤다. 그러나 분석의 편의를 위해, 그 경로를 다섯 단계로 단순화할 수 있다. 즉, (i) 기초적 지식의 생성, (ii) 발명(invention)으로의 결정, (iii) 혁신(innovation)으로의 상업화, (iv) 시장에서의 확산(diffusion), (v) 경제적 가치(GDP·고용·소비자잉여)로의 실현이 그것이다.",
    "이 가치사슬의 각 단계는 서로 다른 시간 척도와 행위자, 그리고 성공률을 갖는다. 기초적 지식의 생성은 대학·정부연구소·기업기초연구실에서 수십 년에 걸쳐 일어나며, 한 편의 논문이 ‘산업적 영향’을 미치기까지 평균 15~20년이 걸린다는 추정이 일반적이다(Mansfield, 1991, 1998). 발명은 특허로 결정되는데, 한 분야의 모든 특허 중 시장에서 상업적 성공을 거두는 비율은 5% 미만이다(Scherer & Harhoff, 2000). 혁신과 확산의 단계에서는 ‘죽음의 계곡(valley of death)’으로 불리는 단계적 실패가 빈번하며, 그 통과의 핵심은 자본의 조달과 흡수역량의 결합이다.",
    "그렇다면 ‘선형 모델(linear model)’—과학 → 기술 → 시장의 일방향 흐름—은 현실의 어떤 부분을 포착하고 어떤 부분을 놓치는가? 첫째, 선형 모델은 ‘출발점’의 단순성에서 강점이 있다. 둘째, 그러나 실제 혁신 사례를 자세히 들여다보면, 시장의 요구(market pull)와 기술의 추진(technology push)이 양방향으로 상호작용한다. 셋째, 자주 인용되는 Kline & Rosenberg(1986)의 ‘연쇄–결합 모형(chain-linked model)’은 이 양방향성을 강조하며, 과학과 기술 사이의 ‘지속적 피드백 루프’를 명시적으로 도입한다.",
    "본 논문은 ‘선형 모델’의 단순성과 ‘연쇄–결합 모형’의 풍부함을 함께 활용한다. 즉, 가치 흐름의 거시적 윤곽을 설명할 때는 5단계 선형을 차용하되, 각 단계의 내부 동학을 분석할 때는 양방향 피드백 구조를 명시한다.",
]
for p in ch3_1:
    add_paragraph(doc, p)

add_figure_caption(doc, "3.1", "과학–기술–혁신–가치의 가치사슬")

ch3_1_more = [
    "[그림 3.1]은 본 논문이 채택한 가치사슬의 개념도이다. 다섯 개의 노드(과학–기술–혁신–확산–가치)는 양방향 화살표로 연결되며, 각 화살표 위에는 통과를 매개하는 ‘제도·조직·인지’ 요인들이 라벨로 표시된다. 본 절 이후 3.2~3.6절은 각 화살표의 주요 매개 요인을 차례로 다룬다.",
]
for p in ch3_1_more:
    add_paragraph(doc, p)

add_heading(doc, "3.2  기술이전과 사업화의 미시 메커니즘", level=2)
ch3_2 = [
    "‘기술이전(technology transfer)’이란 한 조직에서 산출된 기술적 지식이 다른 조직 또는 시장으로 이동하여 새로운 가치를 창출하는 과정이다. 그 채널은 매우 다양하다. 첫째, 라이선싱(licensing)—특허·노하우·저작권의 사용권 부여이다. 둘째, 분사창업(spin-off)—연구자 자신이 발명을 사업화하기 위해 회사를 설립하는 경로이다. 셋째, 컨설팅·자문—연구자가 기업의 특정 문제에 자신의 지식을 적용하는 형태이다. 넷째, 공동연구(joint R&D)—대학·연구소·기업이 자원을 공유하며 새로운 지식을 함께 생성하는 형태이다. 다섯째, 인력 이동(personnel mobility)—연구자 본인이 다른 조직으로 옮기며 ‘몸으로 운반하는’ 암묵지(tacit knowledge)의 이동이다.",
    "각 채널의 효율성은 기술의 성격에 따라 다르다. 코드화(codification) 정도가 높은 명시지(explicit knowledge)는 라이선싱과 매뉴얼화된 절차로 이전이 가능하지만, 암묵지의 비중이 큰 경우 인력 이동이나 장기 공동연구가 사실상 유일한 효과적 채널이다. Polanyi(1967)와 Nonaka & Takeuchi(1995)의 ‘지식 변환’ 이론은 명시지–암묵지 변환의 어려움을 강조하였다.",
    "기술이전의 ‘제도적 인프라’로서 가장 중요한 장치는 대학·연구소의 기술이전조직(Technology Transfer Office, TTO)이다. 미국의 경우 1980년 Bayh-Dole법 이후 대학들이 TTO를 본격적으로 설립하기 시작하였고, 1996~2020년 사이 미국 대학들은 누적 18,000개 이상의 스핀오프 창업과 141,000여 건의 특허를 산출하였다(AUTM Survey). 한국의 경우 1996년 ‘기술이전촉진법’ 제정 이후 대학·출연(연)이 단계적으로 TLO(Technology Licensing Office)를 확장하였고, 2000년대 중반부터는 기술지주회사 및 사업화 자회사를 통해 후속 자본 조달을 지원하는 모델로 진화하였다.",
]
for p in ch3_2:
    add_paragraph(doc, p)

add_table_caption(doc, "3.1", "기술이전 채널과 그 특성")
add_table(
    doc,
    header=["채널", "주요 자산", "암묵지 의존도", "거래비용", "대표 사례"],
    rows=[
        ["라이선싱", "특허·저작권", "낮음", "중", "퀄컴의 CDMA 표준 라이선싱"],
        ["분사창업", "연구자·특허·시설", "중-높음", "고", "구글(스탠퍼드 데이터마이닝 연구실)"],
        ["컨설팅", "전문가의 시간", "높음", "낮음", "교수의 산업 자문"],
        ["공동연구", "사람·시설·데이터", "중-높음", "중-고", "ITRI·삼성 공동개발 사업"],
        ["인력 이동", "암묵지·네트워크", "매우 높음", "거래비용 자체 미정",
         "벤처에 합류하는 박사후연구원"],
    ],
    col_widths_cm=[2.5, 3.0, 2.5, 2.0, 5.0],
)

add_heading(doc, "3.3  트리플 헬릭스: 대학–산업–정부의 상호작용", level=2)
ch3_3 = [
    "에츠코위츠(Henry Etzkowitz)와 라이데스도르프(Loet Leydesdorff)가 1990년대 중반에 정식화한 ‘트리플 헬릭스(triple helix)’ 모형은 지식기반경제(knowledge-based economy)에서 혁신의 본질을 ‘대학–산업–정부’ 세 주체의 ‘상호 침투(mutual interpenetration)’로 본다. 이 모형은 세 단계의 진화로 묘사된다. (1) 1단계는 각 주체가 분리된 영역을 가진 ‘진보주의적(statist)’ 모형이다. (2) 2단계는 정부가 산업과 대학을 ‘분리(laissez-faire)’시키되 조정하는 모형이다. (3) 3단계는 세 주체가 서로의 영역으로 침투하며 상호 보완·대체하는 ‘하이브리드 조직(hybrid organizations)’의 시대이다.",
    "트리플 헬릭스 3단계의 가장 가시적 산물은 대학발 창업, 산업–대학 연구 컨소시엄, 그리고 정부가 만든 ‘혁신 중간 조직(intermediary organizations)’이다. 미국 실리콘밸리의 스탠퍼드대학과 주변 벤처캐피털, 영국 케임브리지 사이언스파크, 한국 대덕연구개발특구, 광주의 GIST 첨단융합기술원이 그러한 하이브리드의 사례이다. 본 논문은 GIST 학위논문이라는 형식에 충실하기 위해, 광주를 중심으로 한 광·바이오·AI 클러스터의 사례를 5.3절 ‘한국 ICT 산업의 도약 경로’ 안에서 보다 자세히 다룬다.",
    "트리플 헬릭스 모형의 학술적 기여는 다음과 같다. 첫째, 혁신을 ‘제도 간 경계의 흐려짐’으로 보는 새로운 관점을 제공한다. 둘째, ‘대학의 제3의 임무(third mission)’—교육·연구에 더한 사회·경제적 가치 창출—를 정당화하는 이론적 기반이 된다. 셋째, 정책 입안자에게 ‘세 헬릭스 중 어느 한 부분만 강화하는 정책의 한계’를 강조한다.",
    "비판도 존재한다. 첫째, ‘트리플’이라는 명목이 실제로는 더 다양한 행위자들—시민사회, 자본시장, 미디어, 사용자 공동체—을 단순화한다는 지적이다. Carayannis & Campbell(2009)은 시민사회를 추가한 ‘쿼드러플 헬릭스(quadruple helix)’, 자연환경을 추가한 ‘퀸터플 헬릭스(quintuple helix)’를 제안하였다. 둘째, 트리플 헬릭스가 ‘제도 간 협력’을 다소 낭만화한다는 지적이다. 실제 협력은 종종 ‘이해 충돌’과 ‘권력 비대칭’을 동반하며, 이를 명시적으로 다루지 않는 모형의 한계가 있다.",
]
for p in ch3_3:
    add_paragraph(doc, p)

add_figure_caption(doc, "3.2", "트리플 헬릭스 모형의 진화")

add_heading(doc, "3.4  흡수역량과 조직 학습", level=2)
ch3_4 = [
    "코헨과 레빈탈(Cohen & Levinthal, 1990)이 제안한 ‘흡수역량(absorptive capacity)’ 개념은 기업이 ‘외부의 새로운 지식을 인지·동화·적용할 수 있는 능력’을 가리킨다. 그들의 핵심 통찰은, 흡수역량 자체가 과거의 R&D 투자와 관련 지식의 누적에 의해 형성된다는 점이다. 즉, ‘배우기 위해서는 먼저 배워야 한다(learning to learn)’. 따라서 R&D는 단순히 새로운 지식을 ‘생산’할 뿐 아니라 외부 지식을 ‘흡수’할 능력 자체를 길러내는 이중 역할을 수행한다.",
    "흡수역량의 구성 요소는 통상 (1) 사전 지식 기반의 폭과 깊이, (2) 조직 내 의사소통의 빈도와 질, (3) 다양한 부서 간 정보 통합 메커니즘, (4) 외부 네트워크와의 연결성, (5) 학습 지향 조직문화로 정리된다. Zahra & George(2002)는 이를 ‘잠재적 흡수역량(potential absorptive capacity, PACAP)’—획득과 동화—과 ‘실현된 흡수역량(realized absorptive capacity, RACAP)’—변환과 활용—으로 구분하였다.",
    "흡수역량 개념은 본 논문의 핵심 주제와 깊이 연결된다. 같은 양의 ‘외부 지식 공급(예: 대학 논문, 정부 R&D 보고서)’이 주어졌을 때, 그것을 기업의 가치로 전환하는 효율은 흡수역량에 의해 크게 달라진다. 한국 사례에서, 1980~90년대 삼성·LG·현대의 흡수역량 누적(해외 라이선싱, 자체 R&D, 박사급 인력의 대규모 채용)이 1990년대 후반의 메모리 반도체·디스플레이 추격을 가능하게 한 결정적 요인이었다는 점은 다수의 연구가 일관되게 지적한다.",
]
for p in ch3_4:
    add_paragraph(doc, p)

add_table_caption(doc, "3.2", "흡수역량의 구성 요소(Cohen & Levinthal, 1990 기반)")
add_table(
    doc,
    header=["구성 요소", "조작적 정의", "측정 변수"],
    rows=[
        ["사전 지식 기반", "기업이 보유한 관련 지식의 폭/깊이",
         "R&D 인력 비중, 누적 R&D 스톡"],
        ["내부 소통", "부서 간/직군 간 정보 흐름",
         "교차기능팀 비율, 회의 빈도"],
        ["외부 연결성", "대학·연구소·고객·공급자 네트워크",
         "공동연구 건수, 공동특허"],
        ["학습 문화", "실패 허용, 실험 장려",
         "혁신 관련 KPI, 사내 학습시간"],
        ["변환 능력", "기존 지식과 신지식의 통합 능력",
         "특허 인용 다양성"],
        ["활용 능력", "흡수된 지식의 상업적 적용",
         "신제품 매출 비중"],
    ],
    col_widths_cm=[3.0, 5.5, 5.5],
)

add_heading(doc, "3.5  네트워크 효과, 플랫폼, 양면시장", level=2)
ch3_5 = [
    "디지털 경제에서 가치 창출의 동학은 전통 산업과 다르다. Rochet & Tirole(2003)이 정식화한 ‘양면시장(two-sided market)’ 이론은 플랫폼이 두 개 이상의 사용자 집단 간 상호작용을 매개하며, 한 집단의 사용자가 증가할수록 다른 집단에 대한 가치가 커지는 ‘교차 네트워크 효과(cross-side network effect)’를 가진다고 본다. 이러한 양면시장에서는 가격이 ‘평균 비용 + 한계비용’ 원리로 결정되지 않으며, 한 쪽 집단에는 보조금(또는 무료 서비스)이, 다른 쪽 집단에는 프리미엄 가격이 부과되는 구조가 일반화된다.",
    "양면시장 동학은 본 논문의 핵심 질문에 두 가지 측면에서 영향을 미친다. 첫째, ‘과학기술이 가치가 되는 경로’가 더 이상 ‘기술 → 상품 → 가격’이라는 단순 형태가 아니다. 데이터·소프트웨어·알고리즘이라는 비경합재가 ‘플랫폼’이라는 양면 구조 안에서 가치로 실현되며, 그 가치의 측정은 GDP에서 종종 잡히지 않는다(Brynjolfsson et al., 2019). 둘째, 플랫폼은 ‘승자독식(winner-take-most)’ 동학을 가지므로, 경쟁정책·반독점정책과의 긴장이 첨예해진다. 본 논문 제6.5절의 ‘책임 있는 혁신’ 논의에서 이 문제를 다룬다.",
]
for p in ch3_5:
    add_paragraph(doc, p)

add_figure_caption(doc, "3.3", "플랫폼 양면시장 구조")

add_heading(doc, "3.6  재조합적 혁신과 일반목적기술(GPT)", level=2)
ch3_6 = [
    "Weitzman(1998)의 ‘재조합적 성장(recombinant growth)’ 모형은 새로운 지식이 기존 지식 요소들의 ‘조합’으로부터 산출된다고 본다. 이 시각에서 보면, 지식 스톡이 클수록 가능한 조합의 수가 폭증하므로, 장기적으로 지식의 한계생산성이 체감하지 않거나 오히려 증가할 수 있다. Arthur(2009)는 『기술의 본질(The Nature of Technology)』에서 이 통찰을 일반화하여, 모든 기술은 기존 기술 요소들의 ‘반복적 재구성’의 산물이라고 주장하였다.",
    "이러한 재조합적 관점은 ‘일반목적기술(general-purpose technology, GPT)’ 개념과 결합한다. Bresnahan & Trajtenberg(1995)는 증기기관, 전기, 정보기술처럼 (i) 광범위한 응용 가능성, (ii) 지속적 개선, (iii) 보완적 혁신을 유발하는 능력을 가진 기술을 GPT로 정의하였다. GPT는 단기적으로는 생산성 역설(productivity paradox)—기술 도입 직후 생산성이 일시적으로 둔화되는 현상—을 보이지만, 장기적으로 누적적 성장 효과가 매우 크다(David, 1990).",
    "2020년대 중반에 등장한 대규모 언어 모델(LLM), 그리고 이를 토대로 한 생성형 AI는 그 자체가 GPT의 가장 강력한 후보로 평가받고 있다(Brynjolfsson, Rock & Syverson, 2017, ‘AI as a GPT’). 본 논문 제5.5절에서 이 사례를 보다 자세히 다룬다.",
]
for p in ch3_6:
    add_paragraph(doc, p)

# ===================== 제4장 =====================
add_heading(doc, "제4장  경제적 가치의 측정", level=1, page_break=True)

add_heading(doc, "4.1  R&D 통계의 국제 표준: Frascati Manual", level=2)
ch4_1 = [
    "OECD가 1963년 처음 발간하고 2015년 7차 개정한 『Frascati Manual』은 R&D 통계의 국제 표준이다. 그 정의에 따르면, R&D는 ‘인간·문화·사회에 대한 지식을 포함한 지식의 스톡을 늘리기 위해, 그리고 이용 가능한 지식을 새로운 응용으로 고안하기 위해 수행되는 체계적인 창조적 작업’이다. R&D 활동이 갖춰야 할 다섯 가지 기준은 (i) 신규성(novel), (ii) 창의성(creative), (iii) 불확실성(uncertain), (iv) 체계성(systematic), (v) 이전가능성·재현가능성(transferable, reproducible)이다.",
    "Frascati Manual은 R&D를 세 가지 유형—기초연구(basic research), 응용연구(applied research), 개발연구(experimental development)—으로 분류한다. 기초연구는 ‘특정한 응용을 일차적으로 염두에 두지 않는, 새로운 지식의 획득을 위한 실험적·이론적 작업’이며, 응용연구는 ‘새로운 지식의 획득을 일차적 목표로 하되 특정한 실용적 목적을 가진 작업’이다. 개발연구는 ‘기존 지식을 활용하여 새로운 재료·제품·공정·시스템·서비스를 생산하거나 기존 것을 실질적으로 개선하기 위한 체계적 작업’이다.",
]
for p in ch4_1:
    add_paragraph(doc, p)

add_table_caption(doc, "4.1", "Frascati Manual의 R&D 분류(2015 개정)")
add_table(
    doc,
    header=["유형", "정의", "예시", "주 수행자"],
    rows=[
        ["기초연구", "응용 목적 없는 새로운 지식 획득",
         "양자장 이론, 단백질 구조 규명", "대학·정부연구소"],
        ["응용연구", "특정 응용을 염두에 둔 새로운 지식 획득",
         "암 치료제 표적 발굴", "대학·정부연구소·기업 중앙연구소"],
        ["개발연구", "기존 지식의 응용으로 새 제품·공정 산출",
         "신차 모델 개발, 양산 공정 최적화", "기업 사업부"],
    ],
    col_widths_cm=[2.5, 5.0, 5.0, 3.5],
)

ch4_1_more = [
    "Frascati Manual은 또한 ‘수행자별/재원자별(performer/funder)’ R&D를 구분한다. 수행자는 R&D를 실제로 수행하는 부문이며, 재원자는 그 비용을 지불하는 부문이다. 통상의 행렬은 (수행자) 기업·정부·고등교육·민간비영리 ×(재원자) 기업·정부·고등교육·민간비영리·해외의 8x5 매트릭스이며, 이 매트릭스의 셀별 합계가 GERD(Gross Domestic Expenditure on R&D)의 핵심 분해이다.",
    "측정의 어려움은 다양하다. 첫째, ‘R&D와 비R&D의 경계’가 모호한 경우가 많다. 예컨대 소프트웨어 개발의 어떤 부분이 R&D인지에 대한 가이드라인은 OECD가 2025년 별도 보고서로 보강하고 있다. 둘째, 다국적 기업의 R&D 활동은 본사와 자회사 간 회계 처리에 따라 측정값이 크게 달라진다. 셋째, ‘무엇이 R&D 수행자인가’의 정의가 디지털 시대에 갈수록 모호해진다(예: 오픈소스 커뮤니티).",
]
for p in ch4_1_more:
    add_paragraph(doc, p)

add_heading(doc, "4.2  혁신활동의 측정: Oslo Manual", level=2)
ch4_2 = [
    "Frascati Manual이 R&D 자체의 ‘투입’을 측정한다면, Oslo Manual은 ‘혁신 산출’과 ‘혁신 활동(innovation activities)’을 측정한다. 1992년 1판 이후 2018년 4판이 발간되었다. 4판의 가장 큰 변화는 혁신의 정의를 단순화·확장한 것이다. 새 정의에 따르면 ‘혁신은 단위(기업)의 기존 제품/공정/마케팅/조직과 의미 있게 다른 새로운 또는 개선된 제품/공정이 사용자에게 가용해지거나 단위에 의해 사용에 도입되는 것’이다. 또한 ‘혁신 활동’은 ‘혁신을 결과로 산출하기 위해 의도된 단위의 모든 발전·재무·상업 활동’으로 확장 정의된다.",
]
for p in ch4_2:
    add_paragraph(doc, p)

add_table_caption(doc, "4.2", "Oslo Manual에 의한 혁신의 4유형")
add_table(
    doc,
    header=["유형", "정의", "예시"],
    rows=[
        ["제품 혁신", "새로운 또는 개선된 제품/서비스의 도입",
         "신약, 신형 스마트폰"],
        ["공정 혁신", "새로운 또는 개선된 생산·전달 공정의 도입",
         "자동화 공정, 클라우드 전환"],
        ["마케팅 혁신", "새로운 마케팅 방법의 도입",
         "온라인 구독 모델"],
        ["조직 혁신", "새로운 조직 구조·관행의 도입",
         "애자일 조직 전환"],
    ],
    col_widths_cm=[2.5, 5.5, 6.0],
)

ch4_2_more = [
    "한국통계청은 2002년부터 ‘한국 혁신조사(KIS)’를 격년으로 수행하여 Oslo Manual에 부합하는 자료를 산출한다. 이 자료는 KISTEP과 STEPI의 정책 분석, 그리고 본 논문 제6장의 분석에 활용된다.",
]
for p in ch4_2_more:
    add_paragraph(doc, p)

add_heading(doc, "4.3  특허·인용·논문계량의 지표학", level=2)
ch4_3 = [
    "특허(patent)는 발명의 가장 가시적인 결과이며, 100년이 넘는 기간에 걸쳐 정형화된 데이터가 축적되어 있어 혁신연구의 ‘준(準)자연실험’ 자원으로 자주 활용된다. 특허 지표의 주요 종류는 (i) 단순 출원·등록 수, (ii) 특허 가족(family) 크기, (iii) 청구항(claim) 수, (iv) 후행 인용(forward citation) 수, (v) 기술 분류의 다양성이다. 이 중 후행 인용 수는 한 특허가 후속 발명에 미친 영향력의 대용변수로 가장 자주 사용된다(Trajtenberg, 1990).",
    "그러나 특허 지표에는 한계도 분명하다. 첫째, ‘발명의 일부’만이 특허화된다. 영업비밀로 보호되는 비중이 상당하며, 특히 소프트웨어 분야의 특허화 관행은 국가별로 큰 차이를 보인다. 둘째, 특허의 ‘질’은 양적 분포가 매우 비대칭적이다. 한 특허의 가치는 후행 인용 분포와 매우 유사한 형태를 보이며, 상위 1%의 특허가 전체 가치의 절반 이상을 설명한다는 추정이 일반적이다(Hall, Jaffe & Trajtenberg, 2005). 셋째, 인용은 부분적으로 ‘대기업의 변호인 관행’의 산물이며, 발명의 실질적 영향과 항상 일치하지는 않는다.",
    "논문계량(bibliometrics)은 과학적 산출의 측정 도구이다. SCI/SSCI 등 인용지수 데이터베이스를 활용한 논문 수·인용 수·인용 영향력(impact factor)·H 지수 등은 연구자·연구실·국가 차원에서 표준적으로 사용된다. 그러나 ‘과학적 산출이 산업적 가치로 전환되는 효율’을 측정하기에는 그 자체로 한계가 있다. 본 논문은 논문계량을 ‘투입 측면의 변형 변수’로 보고, ‘산업적 영향’의 측정에는 특허–논문 인용 매트릭스(Patent-Paper Citation Matrix), 산업의 특허화 강도, 그리고 신생 산업의 시장 형성 등을 결합한 다층 접근을 채택한다.",
]
for p in ch4_3:
    add_paragraph(doc, p)

add_heading(doc, "4.4  GDP 산정에서의 R&D 자본화", level=2)
ch4_4 = [
    "전통적 국민계정(System of National Accounts, SNA)에서 R&D 지출은 중간소비로 분류되어 GDP에 ‘이중 산입’되지 않았다. 그러나 2008 SNA 개정(국제적으로는 2014년경 적용)부터 R&D 지출은 ‘고정자본 형성(gross fixed capital formation, GFCF)’의 한 항목으로 자본화되었다. 이는 R&D가 단순한 비용이 아니라 ‘장기적으로 생산성에 기여하는 자산’이라는 이론적 인식의 결과이다.",
    "R&D 자본화의 회계적 함의는 다음과 같다. 첫째, 모든 국가의 GDP 수준이 R&D 강도에 비례하여 상향 조정된다. 미국의 경우 2013년 BEA의 재계산으로 GDP가 약 2.5% 상향되었다. 둘째, R&D 자본의 감가상각률을 무엇으로 가정하느냐가 결과에 영향을 미친다. OECD는 산업별로 10~25%의 감가상각률을 권장하며, 한국도 이를 따른다. 셋째, R&D 자본 스톡과 다른 자산(설비, 소프트웨어, 데이터)과의 상호작용을 분석할 수 있는 자료가 새롭게 확보된다.",
    "본 논문은 4.4절에서 한국과 OECD 주요국의 R&D 강도(GERD/GDP)와 R&D 자본 스톡의 누적 추이를 비교한다. 한국은 GERD/GDP에서 2023년 기준 4.96%로 OECD 1위 수준이지만, R&D 자본 스톡의 ‘질(기초연구 비중·해외 협력 비중·대학 비중)’에서는 미국·EU에 비해 구조적 약점을 안고 있다는 점이 드러난다.",
]
for p in ch4_4:
    add_paragraph(doc, p)

add_table_caption(doc, "4.3", "주요국 GERD/GDP 비중(2022, 단위: %)")
add_table(
    doc,
    header=["국가", "GERD/GDP(%)", "기업 비중", "정부·대학 비중"],
    rows=[
        ["대한민국", "5.21", "78.5", "21.5"],
        ["이스라엘", "5.56", "89.2", "10.8"],
        ["미국", "3.45", "72.8", "27.2"],
        ["일본", "3.30", "78.1", "21.9"],
        ["독일", "3.13", "67.0", "33.0"],
        ["중국", "2.43", "76.6", "23.4"],
        ["프랑스", "2.22", "65.5", "34.5"],
        ["영국", "2.91", "68.0", "32.0"],
        ["OECD 평균", "2.73", "72.5", "27.5"],
    ],
    col_widths_cm=[3.5, 3.0, 3.0, 4.5],
)

add_figure_caption(doc, "4.2", "한국 R&D 강도(GDP 대비) 추이")

add_heading(doc, "4.5  무형자산 회계와 기업가치", level=2)
ch4_5 = [
    "기업 차원에서 ‘과학기술의 경제적 가치’는 점차 ‘무형자산(intangible assets)’의 형태로 인식된다. Corrado, Hulten & Sichel(2009)의 선구적 작업은 무형자산을 (i) 컴퓨터화된 정보(computerized information), (ii) 혁신적 자산(innovative property, R&D 포함), (iii) 경제적 역량(economic competencies, 브랜드·인적자본·조직자본)으로 분류하였다. 이들은 미국에서 무형자산 투자가 1990년대 말에 이미 유형자산 투자를 추월하였으며, 무형자산 자본 스톡이 노동생산성 증가의 27%를 설명한다고 보고하였다.",
    "회계 표준은 이 변화를 부분적으로만 반영하고 있다. IFRS와 K-IFRS는 ‘식별 가능하고 미래 경제적 편익을 통제할 수 있으며 신뢰성 있게 측정 가능한 자원’만을 무형자산으로 인식하므로, 자체 개발 R&D의 인식 요건은 매우 까다롭다. 결과적으로 재무제표에 기재된 무형자산 가치는 시장가치의 일부분일 뿐이며, 시가총액과 장부가의 격차가 무형자산의 ‘회계 누락분(off-balance-sheet intangibles)’을 반영한다는 해석이 일반적이다(Lev & Gu, 2016).",
    "이 측정의 ‘구멍’은 본 논문 제6장의 정책 논의에 직접 영향을 미친다. 무형자산의 비중이 큰 기업·산업에 대해서는 전통적인 ‘유형자산 투자세액공제’ 같은 정책 도구가 효과를 보장하기 어렵다. 따라서 R&D 세액공제, 무형자산 상각 가속화, 그리고 데이터·소프트웨어 자본화에 대한 별도의 정책 처방이 필요하다.",
]
for p in ch4_5:
    add_paragraph(doc, p)

add_heading(doc, "4.6  사회적 수익률 vs. 사적 수익률", level=2)
ch4_6 = [
    "R&D의 ‘사회적 수익률(social rate of return)’과 ‘사적 수익률(private rate of return)’의 차이는 R&D 정책 설계의 핵심이다. 사적 수익률은 R&D를 수행한 기업이 직접 회수하는 이익이며, 사회적 수익률은 그에 더해 ‘파급효과(spillover)’—다른 기업·산업·소비자에게 전이되는 가치—를 포함한 전체 수익이다. 메타분석들은 R&D의 사회적 수익률이 사적 수익률보다 평균적으로 2~4배 크다고 일관되게 보고한다(Hall, Mairesse & Mohnen, 2010).",
    "이 격차의 함의는 명료하다. 사적 부문 단독으로는 사회적으로 최적인 양보다 적게 R&D에 투자할 것이므로, 공공의 R&D 보조·공공연구 직접수행이 사회후생을 증가시킨다. 그러나 ‘얼마나’ 보조하고 ‘어떤’ 분야에 보조할 것인가는 정밀한 추정과 정치적 합의의 문제이다. 메타분석의 또 다른 결과로, 기초연구의 사회적 수익률이 응용연구보다 통계적으로 유의하게 높다는 추정이 다수 보고된다(예: Jones & Williams, 1998).",
]
for p in ch4_6:
    add_paragraph(doc, p)

add_table_caption(doc, "4.4", "R&D 사회적 수익률 추정치 메타분석(주요 연구)")
add_table(
    doc,
    header=["연구", "데이터/방법", "사적 수익률", "사회적 수익률"],
    rows=[
        ["Mansfield(1977)", "산업별 사례 17건", "약 25%", "약 56%"],
        ["Griliches(1992) 요약", "다수 연구 메타", "20~30%", "50% 이상"],
        ["Jones & Williams(1998)", "거시 캘리브레이션", "—", "30~50%"],
        ["Hall, Mairesse & Mohnen(2010)", "메타분석",
         "20~30%", "30~100%"],
        ["Sveikauskas(2007)", "미국 산업 패널", "약 25%", "약 70%"],
        ["Eberhardt et al.(2013)", "OECD 패널", "—", "약 35%(평균)"],
    ],
    col_widths_cm=[4.5, 4.0, 2.5, 3.0],
)

# ===================== 제5장 =====================
add_heading(doc, "제5장  사   례   연   구", level=1, page_break=True)

add_heading(doc, "5.1  미국 Bayh-Dole 체제와 대학 기술이전", level=2)
ch5_1 = [
    "1980년 12월 12일 미국 의회를 통과한 「특허·상표법 개정법」, 통칭 Bayh-Dole 법은 미국의 대학·중소기업·비영리연구기관이 연방정부의 자금을 사용하여 산출한 발명에 대한 소유권을 보유할 수 있도록 허용하였다. 이전까지는 연방정부 자금이 투입된 발명은 정부 소유였으며, 그 라이선싱 절차의 비효율로 인해 ‘서랍에 잠자는 특허(patents in the drawer)’가 누적되어 있었다. Bayh-Dole 법은 이 비효율을 일시에 해소하였고, 그 결과는 극적이었다.",
    "AUTM(Association of University Technology Managers)이 매년 발간하는 라이선싱 활동조사(STATT) 자료에 따르면, 1996년부터 2020년까지 25년간 미국 대학에서 발생한 누적 성과는 다음과 같다. 첫째, 발명 공시(invention disclosure) 약 554,000건. 둘째, 미국 특허 등록 약 141,000건. 셋째, 누적 스핀오프 창업 약 18,000개. 넷째, 미국 산업총산출에 대한 누적 기여 약 1.9조 달러. 다섯째, 누적 고용 기여 약 690만 명.",
]
for p in ch5_1:
    add_paragraph(doc, p)

add_table_caption(doc, "5.1", "Bayh-Dole 이후 미국 대학 기술이전 성과(1996–2020 누적)")
add_table(
    doc,
    header=["지표", "누적 값", "참고"],
    rows=[
        ["발명 공시(disclosure)", "약 554,000건", "AUTM STATT"],
        ["미국 특허 등록", "약 141,000건", "USPTO"],
        ["라이선싱 계약 체결", "약 105,000건", "AUTM"],
        ["대학 스핀오프 창업", "약 18,000개", "AUTM"],
        ["산업총산출 기여", "약 1.9조 USD", "BIO·ITIF 추정"],
        ["누적 고용 기여", "약 690만 명", "ITIF(2025)"],
        ["연간 신제품 출시", "평균 2건/일", "AUTM"],
    ],
    col_widths_cm=[5.5, 4.0, 4.5],
)

ch5_1_more = [
    "Bayh-Dole이 ‘성공의 모범’으로 자주 인용되지만, 비판도 적지 않다. 첫째, ‘무한정한 라이선싱 권한’이 기초연구의 공공성(특히 의약품 접근성)을 침식한다는 우려가 있다. 둘째, 대학 TTO의 ‘라이선싱 수수료 극대화’ 행태가 단기·소형 라이선싱에 치우치게 한다는 비판이 있다. 셋째, ‘승자독식’의 현실—극소수의 ‘대박 특허’가 전체 라이선싱 수입의 대부분을 설명—이 대학 간 격차를 누적적으로 강화한다는 지적이 있다.",
    "Bayh-Dole의 한국적 함의는 1996년 「기술이전촉진법」, 2000년 「기술이전 및 사업화 촉진에 관한 법률」로 이어진 일련의 입법에서 확인된다. 한국은 미국 모델을 부분적으로 차용하면서, 한국 특유의 거버넌스—정부출연(연)의 비중이 크고, 대기업 중심의 자체 R&D가 강한 구조—에 맞게 변형하였다. 본 논문 5.3·5.4·6.2절은 이 한국적 변형을 보다 자세히 다룬다.",
]
for p in ch5_1_more:
    add_paragraph(doc, p)

add_figure_caption(doc, "5.1", "미국 대학 기술이전의 누적 성과(1996–2020)")

add_heading(doc, "5.2  공공 R&D가 잉태한 시장: 인터넷·GPS·터치스크린·mRNA",
            level=2)
ch5_2 = [
    "마추카토(2013)가 강조한 ‘기업가형 국가’의 가장 가시적 사례는 다음의 네 가지이다.",
    "첫째, 인터넷이다. 1960년대 후반 미국 국방고등연구계획국(DARPA)이 자금을 댄 ARPANET이 그 뿌리이며, 1970~80년대 TCP/IP 프로토콜의 표준화, 1980년대 NSFNET을 통한 학술망 확장, 1990년대 후반 민간에의 상업 개방을 거치며 오늘날의 글로벌 인터넷이 성립하였다. 핵심 기술 요소—패킷 교환, 라우팅 알고리즘, 도메인 네임 시스템—는 거의 모두 공공 자금이 투입된 연구의 산물이다.",
    "둘째, GPS이다. 1973년 미 국방부의 NAVSTAR GPS 프로그램으로 시작된 위성항법시스템은 1995년 완전 운용능력(FOC)에 도달하였고, 2000년 ‘선택적 가용성(SA)’ 해제로 민간 정밀도가 크게 향상되었다. GPS는 오늘날 항공·해운·자율주행·정밀농업·금융거래 시각동기 등 거의 모든 디지털 산업의 ‘보이지 않는 인프라’로 작동한다. 그 누적 경제적 가치는 연간 1조 달러를 넘는 것으로 추정된다(RTI, 2019).",
    "셋째, 터치스크린이다. 1965년 영국 왕립 레이더 시설(Royal Radar Establishment)에서 처음 정전식 터치스크린이 개발되었고, 1970~80년대 미 국방부 자금과 CERN 등 공공 부문에서 다양한 형태가 개량되었다. 1980년대 후반 미 국립과학재단(NSF)의 자금이 투입된 델라웨어대학의 연구는 후일 애플 아이폰의 멀티터치 기술의 원형이 되었다.",
    "넷째, mRNA 백신이다. 1980~90년대 펜실베이니아대학의 카탈린 카리코(Katalin Karikó)와 드루 와이스만(Drew Weissman)의 연구가 그 토대이며, 1990년대부터 미국 NIH의 지속적 자금 지원, 2000년대 NIAID의 RNA 변형 기술 지원, 2020년 BARDA·OWS(Operation Warp Speed)의 본격적 자본 투입을 거쳐 모더나와 BioNTech의 백신으로 결실을 맺었다. 카리코와 와이스만은 2023년 노벨생리의학상을 수상하였다.",
    "이 네 사례의 공통점은 다음과 같다. (i) 핵심 기술의 ‘출발점’이 공공 부문이다. (ii) 그 출발점에서 시장 가치 실현까지의 시간은 10~50년이며, 그 사이 다수의 ‘실패’가 있었다. (iii) 시장 가치 실현 단계에서 민간 기업의 역할이 결정적이지만, 그 ‘발명의 위험’의 대부분은 공공이 부담하였다. 본 논문은 이 사례들로부터 ‘공공 위험 부담 vs. 사적 보상 전유’의 비대칭을 어떻게 교정할 것인가라는 정책 질문을 도출하며, 이를 제6장에서 다룬다.",
]
for p in ch5_2:
    add_paragraph(doc, p)

add_heading(doc, "5.3  한국 ICT 산업의 도약 경로", level=2)
ch5_3 = [
    "한국 ICT 산업의 도약은 1980년대 후반의 TDX(전전자교환기) 국산화, 1990년대 중반의 CDMA 디지털 이동통신 세계 최초 상용화, 1990년대 후반의 초고속인터넷 인프라 구축, 2000년대의 메모리·디스플레이 세계 1위, 2010년대의 모바일 단말 1위, 2020년대 초의 5G 세계 최초 상용화로 이어지는 ‘단계적 도약’의 연속이다. 이 도약의 공통 패턴은 ‘공공의 위험 부담 + 대기업의 양산 능력 + 정부의 표준화·인허가·인프라 지원’의 삼각구도이다.",
    "TDX는 한국전자통신연구원(ETRI)과 4개 통신기업의 공동 R&D로 1986년 상용화되었고, 1991년 1,000만 회선 달성으로 한국을 세계 10번째 자체 디지털 교환기 보유국으로 만들었다. CDMA는 1991년 ETRI–퀄컴 기술도입 계약으로 시작하여 1996년 세계 최초 상용화에 성공하였다. 초고속인터넷은 1995년 ‘초고속정보통신망 구축 종합계획’에 따라 정부가 인프라 투자와 시장 형성을 동시에 추진한 결과, 2002년 가구 보급률 70%를 돌파하였다.",
    "이러한 한국 ICT 도약의 ‘제도적 비결’은 무엇인가? 첫째, 정부출연(연)의 ‘응용·개발 중심’ R&D가 빠른 산업화를 가능하게 하였다. 둘째, 대기업의 자체 R&D 역량 누적이 흡수역량의 토대가 되었다. 셋째, 인허가·표준화의 정부 역할이 ‘조정 비용’을 크게 낮추었다. 넷째, 군 경험을 통한 ‘기술 친화적 청년 인력’의 사회 진출이 확산 단계의 노동공급을 채웠다. 다섯째, 학원·대학·대학원의 ‘공학 교육 강조’가 인적자본 공급의 안정성을 유지하였다.",
    "그러나 이 모델의 한계도 점차 드러난다. ‘추격형’ 단계의 정형화된 강점이 ‘선도형’ 단계의 요구—기초연구, 실패 허용 문화, 글로벌 인재 흡수, 창업 생태계—와 잘 맞지 않는다는 비판이 1990년대 후반부터 누적되어 왔다. 본 논문 제6장은 이 ‘추격형 → 선도형’ 전환의 정책 과제를 보다 자세히 다룬다.",
]
for p in ch5_3:
    add_paragraph(doc, p)

add_table_caption(doc, "5.2", "한국 ICT 5대 품목의 세계시장 점유율 변화(%)")
add_table(
    doc,
    header=["품목", "2000년", "2010년", "2020년", "2024년"],
    rows=[
        ["DRAM", "약 32%", "약 49%", "약 71%", "약 70%"],
        ["NAND 플래시", "약 12%", "약 39%", "약 47%", "약 50%"],
        ["대형 OLED 패널", "—", "약 88%", "약 96%", "약 95%"],
        ["스마트폰", "—", "약 31%", "약 22%", "약 20%"],
        ["5G 장비(글로벌 매출)", "—", "—", "약 18%", "약 24%"],
    ],
    col_widths_cm=[5.0, 2.5, 2.5, 2.5, 2.5],
)

add_heading(doc, "5.4  반도체: 자본·기술·정책의 합주", level=2)
ch5_4 = [
    "반도체는 한국 경제의 단일 최대 수출 품목이며, 본 논문의 핵심 사례 산업이다. 1983년 삼성전자의 64K DRAM 자체 개발 성공은 ‘추격형 도전’의 상징이었고, 1992년 16M DRAM 세계 최초 개발, 1994년 256M DRAM 세계 최초 개발은 ‘추격’이 ‘동행’으로 그리고 ‘선도’로 전환하는 분기점이었다. 2024년 기준 한국의 메모리 반도체(DRAM·NAND) 합산 세계 점유율은 약 60% 수준으로, 단일 국가의 단일 산업 점유율로는 세계 경제사에서 드문 사례이다.",
    "이 성공의 기저에는 다음의 요인이 있다. (i) 1970년대 후반~80년대 초의 막대한 ‘선행 투자(forward investment)’의 모험성, (ii) 1990년대 IMF 위기 직전·직후의 ‘구조조정과 집중’, (iii) 2000년대 후반의 ‘적층화·미세화’를 향한 지속적 R&D, (iv) 2010년대의 ‘외주 파운드리 vs. 종합 IDM’ 선택의 결과, (v) 2020년대 미·중 갈등과 ‘반도체 동맹’의 지정학적 환경 변화. 본 논문은 이 다섯 요인을 본 논문 2.5절의 아기온–호윗 모형과 2.6절의 NIS 관점에서 분석한다.",
    "반도체 사례의 가장 큰 정책적 함의는 ‘대규모·장기·고위험’ 투자에 대한 ‘사회적 자본조달 구조’의 중요성이다. 메모리 한 세대의 생산설비 투자에는 30조 원 이상이 소요되고, 그 회수 기간은 5~10년이다. 이 시계는 일반 자본시장의 의사결정 시계를 훨씬 초과하므로, 대주주의 ‘장기적 의지’, 정부의 ‘위험 분담’, 그리고 인력·소재·장비 생태계의 ‘안정성’이 동시에 갖춰져야 한다. 한국은 이 세 조건을 부분적으로 갖췄으나, 2020년대 들어 미국 IRA·CHIPS법, 일본 보조금 정책, 중국 ‘대약진’ 등 외생 충격에 대한 대응 전략이 시험대에 올라 있다.",
]
for p in ch5_4:
    add_paragraph(doc, p)

add_figure_caption(doc, "5.2", "한국 ICT·반도체 수출 비중 변화")

add_heading(doc, "5.5  AI/LLM과 지능의 산업화", level=2)
ch5_5 = [
    "2017년 Google의 「Attention Is All You Need」 논문이 발표한 트랜스포머(Transformer) 아키텍처는 자연어처리 분야의 패러다임을 바꾸었고, 그 5년 후인 2022년 OpenAI의 ChatGPT 공개는 ‘LLM 시대’의 상업적 개막을 알렸다. 2024~2025년 사이 글로벌 LLM 시장은 폭발적 성장세를 보였고, 2025년 한국은 ‘AI 컴퓨팅 인프라·모델·응용’의 3축에서 모두 의미 있는 진입을 시도하였다.",
    "AI/LLM 산업의 가치사슬은 (i) 하드웨어(GPU·HBM·전력), (ii) 데이터·파운데이션 모델, (iii) 미들웨어(파인튜닝·RLHF·인프라 SaaS), (iv) 응용(코파일럿·에이전트·도메인 SaaS), (v) 사용자 경험(인터페이스·UX) 다섯 층으로 나뉜다. 한국은 (i)층에서 HBM 메모리 세계 1위, (ii)~(iii)층에서는 네이버 하이퍼클로바X, 카카오 KoGPT, KT 믿:음 등 자체 파운데이션 모델의 일정한 진입, (iv)~(v)층에서는 다수의 스타트업과 SaaS 기업의 활약을 보이고 있다.",
    "AI/LLM 산업이 본 논문에 던지는 가장 큰 질문은 두 가지이다. 첫째, 디지털 비경합재—데이터·파라미터—의 경제적 가치를 어떻게 측정·계상할 것인가? 둘째, ‘승자독식’의 가능성이 높은 산업에서 한국이 어떤 ‘틈새 우위(niche advantage)’를 확보해야 하는가? 본 논문 제6장은 이 두 질문에 대한 시론적 답을 시도한다.",
    "특히 본 논문이 강조하는 바는, AI/LLM이 ‘일반목적기술(GPT)’의 성격을 가지므로 그 가치 실현은 단일 기업이나 단일 산업의 사유물이 아니라 ‘경제 전반의 보완적 혁신’에 의해 결정된다는 점이다. 즉, AI 기술 그 자체가 아니라 ‘AI를 활용한 사업 모델·조직 혁신·인적자본 재배치’가 가치 창출의 본 무대이다. 이 점은 1990년대 ‘솔로우의 컴퓨터 역설(computer paradox)’과 동일한 구조이며, 따라서 ‘생산성 효과의 본격적 발현’까지는 10년 안팎의 시차가 있을 것으로 예상된다(Brynjolfsson et al., 2017; 2024 후속 연구).",
]
for p in ch5_5:
    add_paragraph(doc, p)

# ===================== 제6장 =====================
add_heading(doc, "제6장  한국에의 정책적 함의", level=1, page_break=True)

add_heading(doc, "6.1  한국 R&D 투자의 구조와 변동성", level=2)
ch6_1 = [
    "한국은 GDP 대비 R&D 비중(GERD/GDP) 기준으로 OECD 최상위권을 유지하고 있다. 2022년 5.21%, 2023년 4.96%로 미국·일본·독일·중국을 모두 상회한다. 그러나 이 상위권 위상은 다음의 구조적 약점과 동행한다. 첫째, 기업 R&D 비중이 78%를 넘는 ‘기업 편중성’으로 인해 기초연구의 비중이 OECD 평균에 비해 낮다. 둘째, 정부 R&D 예산의 단기 변동성이 크다. 2024년 정부 R&D 예산이 전년 대비 약 14% 감액된 사건은 연구 현장의 신뢰를 크게 흔들었으며, 다년도 R&D 사업의 ‘끊김’과 인력 이탈을 초래하였다(KISTEP, 2024 보고서).",
]
for p in ch6_1:
    add_paragraph(doc, p)

add_table_caption(doc, "6.1", "한국 정부 R&D 예산 추이(2018–2025, 단위: 조 원)")
add_table(
    doc,
    header=["연도", "정부 R&D 예산", "전년 대비 증감"],
    rows=[
        ["2018", "19.7", "+5.3%"],
        ["2019", "20.5", "+4.1%"],
        ["2020", "24.2", "+18.0%"],
        ["2021", "27.4", "+13.2%"],
        ["2022", "29.8", "+8.8%"],
        ["2023", "31.1", "+4.4%"],
        ["2024", "26.5", "-14.7%"],
        ["2025", "29.7", "+11.9%"],
    ],
    col_widths_cm=[3.0, 5.0, 5.0],
)

ch6_1_more = [
    "정책 제언 1: 다년도 R&D 예산제(multi-year R&D budget)와 ‘안정성 룰’의 도입이다. 정부 R&D 예산의 전년 대비 변동폭을 일정 구간(예: ±5%) 안으로 제한하는 ‘예산 안정성 규준(budget stability rule)’을 도입함으로써, 연구 현장의 예측 가능성과 다년도 사업의 지속성을 확보해야 한다.",
    "정책 제언 2: 기초·응용·개발의 균형 회복이다. 한국의 기초연구 비중은 OECD 평균(약 17%)을 따라잡는 것을 단기 목표로 하되, 그 증가는 정부출연(연)뿐 아니라 대학·기업기초연구실을 통해 다중 채널로 이루어져야 한다.",
]
for p in ch6_1_more:
    add_paragraph(doc, p)

add_heading(doc, "6.2  KISTEP과 평가체계", level=2)
ch6_2 = [
    "한국과학기술기획평가원(KISTEP, Korea Institute of S&T Evaluation and Planning)은 한국의 국가 R&D 시스템에서 ‘기획–배분–평가’의 전 주기를 지원하는 핵심 기관이다. KISTEP의 주요 기능은 (i) 국가 R&D 중장기 계획의 기획 지원, (ii) 정부 R&D 예산의 분석과 조정, (iii) 사업·과제·기관 단위의 평가, (iv) 미래 기술 전망(foresight), (v) 국제 STI 정책 협력이다. 1999년 설립 이후 KISTEP의 평가 보고서는 한국 R&D 정책의 가장 중요한 참고자료 가운데 하나가 되어 왔다.",
    "그럼에도 KISTEP을 비롯한 한국의 R&D 평가체계는 다음의 한계를 안고 있다. 첫째, ‘논문 수·특허 수’ 등 정량지표에 대한 의존이 여전히 크다. 둘째, 사업 단위 평가가 중심이며, ‘체제 효과(system effect)’의 평가가 상대적으로 약하다. 셋째, 평가의 ‘학습 기능(learning function)’보다 ‘책무성 기능(accountability function)’이 우세하며, 그 결과 실패의 학습 가치가 충분히 환원되지 않는다.",
    "정책 제언 3: ‘학습형 평가(learning-oriented evaluation)’의 강화이다. 사업 종료 후의 ‘책무성 평가’와 별도로, 사업 중간의 ‘적응적 학습 평가’를 의무화하고, 그 결과를 후속 사업 설계에 환류하는 거버넌스 절차를 명문화해야 한다.",
    "정책 제언 4: 정성·정량 혼합 지표의 정립이다. 임팩트(impact) 평가에는 정량적 산출 외에 ‘사회적 수용성’, ‘후속 혁신 유발’, ‘인적자본 누적’ 등 정성 지표가 결합되어야 한다.",
]
for p in ch6_2:
    add_paragraph(doc, p)

add_heading(doc, "6.3  임무지향형 혁신정책의 한국적 적용", level=2)
ch6_3 = [
    "마추카토와 OECD가 강조해 온 ‘임무지향형 혁신정책(mission-oriented innovation policy, MOIP)’은 ‘무엇을 해결할 것인가(what to solve)’라는 사회·환경적 미션을 먼저 설정하고, 그 미션 달성을 위한 다부처·다산업·다행위자 협력을 R&D 정책의 중심에 두는 패러다임이다. EU의 「Horizon Europe」의 다섯 미션(암 정복, 기후 적응, 해양·내수, 도시, 토양), 미국 NIH의 「Cancer Moonshot」, 영국의 「Industrial Strategy Challenge Fund」가 그 대표적 사례이다.",
    "한국도 2018년 ‘국가전략혁신성장 8대 동력’, 2020년 ‘디지털 뉴딜·그린 뉴딜’, 2022년 ‘12대 국가전략기술’ 등 일련의 ‘미션 지향’ 프로그램을 시도해 왔다. 그러나 OECD(2025)의 한국 MOIP 검토 보고서는 다음의 한계를 지적한다. (i) 미션 설정의 사회적 합의 과정이 단기적이고 폐쇄적이다. (ii) 다부처 협력의 거버넌스가 명목상에 그치는 경우가 많다. (iii) 미션 달성을 위한 ‘위험 분담 자본(patient capital)’의 공급이 충분하지 않다. (iv) 미션의 ‘진행과 실패’에 대한 투명한 보고 체계가 미흡하다.",
    "정책 제언 5: 미션 거버넌스의 ‘제3섹터 통합’이다. 미션의 설정·이행·평가의 모든 단계에 시민사회·사용자 대표·노동의 참여를 명문화함으로써, 미션의 사회적 정당성과 적응성을 동시에 강화해야 한다.",
    "정책 제언 6: ‘이중 트랙 R&D’의 도입이다. 미션의 단기 응용 트랙과 장기 기초 트랙을 병행 운용하되, 두 트랙의 인력·지식이 양방향으로 흐를 수 있는 ‘회전문(revolving door)’ 제도를 정비해야 한다.",
]
for p in ch6_3:
    add_paragraph(doc, p)

add_table_caption(doc, "6.2", "주요국 임무지향형 혁신정책 비교")
add_table(
    doc,
    header=["국가/연합", "대표 프로그램", "핵심 미션", "거버넌스 특징"],
    rows=[
        ["EU", "Horizon Europe Missions",
         "암·기후·해양·도시·토양", "미션위원회·미션보드의 다층 구조"],
        ["미국", "Cancer Moonshot, ARPA-H",
         "암 사망률 50% 감축 등", "NIH 산하 + ARPA형 도전 자본"],
        ["영국", "Industrial Strategy Challenge Fund",
         "AI·노화·청정성장·미래교통", "UKRI 통합 거버넌스"],
        ["독일", "Hightech Strategy 2025",
         "기후·건강·디지털 주권", "BMBF 주관, 산업 협력 강조"],
        ["일본", "Moonshot R&D Program",
         "9대 도전 목표", "JST/CSTI 합동 운용"],
        ["한국", "12대 국가전략기술 등",
         "반도체·바이오·AI 등", "다부처, 컨트롤타워 정비 중"],
    ],
    col_widths_cm=[3.0, 4.0, 4.0, 4.0],
)

add_figure_caption(doc, "6.1", "임무지향 R&D 거버넌스 개념도")

add_heading(doc, "6.4  글로벌 비교: 미·EU·일·중·이스라엘", level=2)
ch6_4 = [
    "한국의 R&D 정책을 ‘비교 시각’에서 조명하면 그 강점과 약점이 더욱 분명해진다. 미국은 ‘파편화된 거버넌스 + 강한 시장 + 풍부한 위험자본’의 조합으로 ‘예측 불가능한 혁신’의 산실이다. EU는 ‘회원국 다양성 + 미션 지향 + 표준화’의 조합으로 ‘체계적이지만 느린 혁신’의 강점을 보인다. 일본은 ‘대기업 중심 R&D + 장수 기업 + 점진적 개선’의 조합으로 ‘안정적이지만 파괴적 혁신에 약한’ 모델을 형성하였다. 중국은 ‘국가 주도 거대 자본 + 빠른 실행 + 광대한 내수’의 조합으로 ‘추격에서 동행, 일부 분야 선도’로 빠르게 이동하고 있다. 이스라엘은 ‘소형 국가 + 군 R&D + 글로벌 자본’의 조합으로 ‘인구 대비 압도적 혁신 밀도’를 보인다.",
    "한국은 이 다섯 모델 중 일본 모델과 가장 유사한 출발점을 가지지만, 1990년대 이후 ‘대기업 중심 + 정부출연(연) + 표준화의 정부 역할’이라는 변형을 거쳐 왔다. 본 논문이 제안하는 한국형 모델의 방향은 일본의 안정성·이스라엘의 위험자본·EU의 미션지향성을 부분적으로 결합한 ‘하이브리드 한국형 혁신체제’이다.",
]
for p in ch6_4:
    add_paragraph(doc, p)

add_heading(doc, "6.5  책임 있는 혁신과 사회적 수용성", level=2)
ch6_5 = [
    "기술의 경제적 가치는 ‘사회적 수용성(social acceptability)’과 분리될 수 없다. 한국 사회는 1990년대 후반의 황우석 사태, 2000년대 광우병 논쟁, 2010년대 사드 전자파 논쟁, 2020년대 AI 챗봇 ‘이루다’ 사건 등 일련의 ‘기술–사회’ 갈등을 경험하였다. 이 사건들은 ‘기술적 정당성’만으로는 가치 실현이 보장되지 않으며, ‘절차적 정당성’과 ‘분배적 정당성’이 함께 충족되어야 함을 보여 주었다.",
    "유럽연합이 2010년대 중반부터 정착시킨 ‘책임 있는 연구와 혁신(Responsible Research and Innovation, RRI)’ 프레임은 (i) 다양성·포용성(diversity & inclusion), (ii) 예측(anticipation), (iii) 성찰(reflection), (iv) 반응성(responsiveness)의 네 요소를 강조한다. 한국에서도 2021년 「인공지능 윤리기준」, 2024년 「인공지능 기본법」 입법으로 일부 제도화가 진행되고 있다.",
    "정책 제언 7: 책임 있는 혁신의 ‘평가 통합’이다. R&D 사업의 사전·중간·사후 평가에 ‘사회적 수용성·윤리·형평성’ 항목을 의무 포함하고, 그 결과를 사업 설계의 환류 루프 안에 명시적으로 통합해야 한다.",
    "정책 제언 8: ‘기술시민권(technological citizenship)’의 강화이다. 시민이 기술의 사용자이자 위험 부담자라는 인식 하에, 표준 설정·인허가·사후 모니터링 과정에 시민 참여 채널을 제도화해야 한다.",
]
for p in ch6_5:
    add_paragraph(doc, p)

# ===================== 제7장 =====================
add_heading(doc, "제7장  결        론", level=1, page_break=True)

add_heading(doc, "7.1  연구의 요약", level=2)
ch7_1 = [
    "본 논문은 ‘과학기술이 어떻게 경제적 가치가 되는가’라는 질문을 (이론) → (메커니즘) → (측정) → (정책)의 네 단계로 분해하여, 각 단계의 주요 학술적 자원과 정책적 현안을 통합적으로 분석하였다.",
    "이론적 측면에서, 본 논문은 솔로우의 외생적 기술진보 가정에서 출발하여, 로머의 내생적 성장이론, 아기온–호윗의 슘페터적 성장 모형, 넬슨–윈터의 진화경제학, 프리먼·룬드발의 국가혁신체제론, 마추카토의 기업가형 국가론에 이르는 이론적 계보를 검토하였다. 그 결과, ‘과학기술–경제 가치’의 관계를 설명하는 단일 이론은 존재하지 않으며, 거시·미시·시스템의 다층적 관점이 결합되어야 한다는 결론에 이르렀다.",
    "메커니즘 측면에서, 본 논문은 지식–발명–혁신–확산–가치의 가치사슬, 기술이전의 다섯 채널, 트리플 헬릭스의 진화, 흡수역량의 구성, 양면시장과 플랫폼, 재조합적 혁신과 GPT의 개념을 통합하였다. 이를 통해 ‘기술의 가치화’가 단일한 직선이 아니라 ‘다층적·다행위자적·진화적 시스템’임을 보였다.",
    "측정 측면에서, 본 논문은 OECD Frascati·Oslo Manual의 표준, 특허·논문계량의 지표학, R&D 자본화의 회계적 의의, 무형자산 회계의 한계, 사회적·사적 수익률의 격차를 비교하였다. 그 결과, 한국의 R&D 강도는 세계 최상위권이지만, 기초연구 비중·평가체계·무형자산 회계의 측면에서 개선의 여지가 큼을 확인하였다.",
    "정책 측면에서, 본 논문은 한국 R&D 투자의 변동성, KISTEP의 평가체계, 임무지향형 혁신정책의 한국적 적용, 책임 있는 혁신의 제도화 등 네 영역에서 총 8개의 정책 제언을 도출하였다.",
]
for p in ch7_1:
    add_paragraph(doc, p)

add_heading(doc, "7.2  학술적·정책적 시사점", level=2)
ch7_2 = [
    "본 논문의 학술적 시사점은 다음과 같다. 첫째, 그동안 별개 분과로 다루어지던 경제성장 이론, 혁신연구, 기술이전 연구, 측정 방법론, 정책 분석을 단일한 분석 프레임 안에 결합하였다. 둘째, ‘비경합성–배제성’의 지식 속성을 분석의 출발점으로 일관되게 사용함으로써, 다층적 논의의 이론적 일관성을 확보하였다. 셋째, AI·플랫폼 등 디지털 시대의 새로운 가치 창출 양식을 기존 이론의 연속·단절의 시각에서 재정위하였다.",
    "정책적 시사점은 본문 제6장에서 8개의 제언으로 압축되어 있다. 그 핵심을 다시 정리하면 다음과 같다. (i) 다년도 R&D 예산제와 예산 안정성 룰, (ii) 기초·응용·개발의 균형 회복, (iii) 학습형 평가의 강화, (iv) 정성·정량 혼합 지표의 정립, (v) 미션 거버넌스의 제3섹터 통합, (vi) 이중 트랙 R&D와 회전문 인사, (vii) 책임 있는 혁신의 평가 통합, (viii) 기술시민권의 제도화.",
]
for p in ch7_2:
    add_paragraph(doc, p)

add_heading(doc, "7.3  연구의 한계와 향후 과제", level=2)
ch7_3 = [
    "본 연구의 한계는 다음과 같다. 첫째, 이론적 통합 연구의 성격상 새로운 계량경제학적 실증을 직접 수행하지 않았으며, 인용된 추정치들의 식별 전략은 각 원자료의 신뢰성에 의존한다. 둘째, 사례 연구의 선택은 ‘대표 사례 + 한국 사례’ 중심이며, 실패 사례의 체계적 비교가 부족하다. 셋째, 정책 제언의 실현 가능성에 대한 정치경제학적 분석이 본 논문의 범위를 넘는다.",
    "향후 연구 과제는 다음과 같다. 첫째, 한국 R&D 데이터(NTIS)와 기업 패널 자료의 결합을 통한 ‘프로젝트–기업–산업’ 다층 분석. 둘째, AI·플랫폼 산업의 가치 측정 방법론 개발. 셋째, 임무지향형 혁신정책의 한국적 적용에 대한 액션리서치. 넷째, ‘책임 있는 혁신’의 사회적 수용성 측정과 정책 평가의 통합. 다섯째, 광주를 비롯한 지역 혁신 클러스터의 사례연구를 통해 ‘지역 차원의 NIS’의 한국형 모델을 정립하는 것.",
    "본 논문은 시작이자 출발점이다. ‘과학기술이 어떻게 경제적 가치가 되는가’라는 질문은 영원히 새롭게 답해져야 하는 질문이며, 그 답은 시대마다 새로운 형태로 다시 쓰여야 한다. 본 논문이 한국 사회의 ‘과학기술의 경제적 의미’를 묻는 학술적·정책적 대화에 작은 기여가 되기를 기대한다.",
]
for p in ch7_3:
    add_paragraph(doc, p)

# ===================== 참고문헌 =====================
add_heading(doc, "참 고 문 헌", level=1, page_break=True)

references = [
    "Acemoglu, D., Johnson, S., & Robinson, J. A. (2001). The colonial origins of comparative development: An empirical investigation. American Economic Review, 91(5), 1369–1401.",
    "Aghion, P., Bloom, N., Blundell, R., Griffith, R., & Howitt, P. (2005). Competition and innovation: An inverted-U relationship. Quarterly Journal of Economics, 120(2), 701–728.",
    "Aghion, P., & Howitt, P. (1992). A model of growth through creative destruction. Econometrica, 60(2), 323–351.",
    "Arrow, K. J. (1962). The economic implications of learning by doing. Review of Economic Studies, 29(3), 155–173.",
    "Arthur, W. B. (2009). The nature of technology: What it is and how it evolves. Free Press.",
    "Association of University Technology Managers (AUTM). (2021). STATT Survey FY 2020. AUTM.",
    "Bresnahan, T. F., & Trajtenberg, M. (1995). General purpose technologies: ‘Engines of growth’? Journal of Econometrics, 65(1), 83–108.",
    "Brynjolfsson, E., Collis, A., Diewert, W. E., Eggers, F., & Fox, K. J. (2019). GDP-B: Accounting for the value of new and free goods in the digital economy. NBER Working Paper 25695.",
    "Brynjolfsson, E., Rock, D., & Syverson, C. (2017). Artificial intelligence and the modern productivity paradox: A clash of expectations and statistics. NBER Working Paper 24001.",
    "Carayannis, E. G., & Campbell, D. F. J. (2009). ‘Mode 3’ and ‘Quadruple Helix’: Toward a 21st century fractal innovation ecosystem. International Journal of Technology Management, 46(3/4), 201–234.",
    "Coe, D. T., & Helpman, E. (1995). International R&D spillovers. European Economic Review, 39(5), 859–887.",
    "Cohen, W. M., & Levinthal, D. A. (1990). Absorptive capacity: A new perspective on learning and innovation. Administrative Science Quarterly, 35(1), 128–152.",
    "Corrado, C., Hulten, C., & Sichel, D. (2009). Intangible capital and U.S. economic growth. Review of Income and Wealth, 55(3), 661–685.",
    "David, P. A. (1990). The dynamo and the computer: An historical perspective on the modern productivity paradox. American Economic Review, 80(2), 355–361.",
    "Eberhardt, M., Helmers, C., & Strauss, H. (2013). Do spillovers matter when estimating private returns to R&D? Review of Economics and Statistics, 95(2), 436–448.",
    "Easterly, W., & Levine, R. (2001). It’s not factor accumulation: Stylized facts and growth models. World Bank Economic Review, 15(2), 177–219.",
    "Etzkowitz, H., & Leydesdorff, L. (1995). The Triple Helix: University-industry-government relations: A laboratory for knowledge-based economic development. EASST Review, 14(1), 14–19.",
    "Etzkowitz, H. (2003). Innovation in innovation: The triple helix of university–industry–government relations. Social Science Information, 42(3), 293–337.",
    "Freeman, C. (1987). Technology policy and economic performance: Lessons from Japan. Pinter.",
    "Griliches, Z. (1992). The search for R&D spillovers. Scandinavian Journal of Economics, 94, S29–S47.",
    "Griliches, Z. (1995). R&D and productivity: Econometric results and measurement issues. In P. Stoneman (Ed.), Handbook of the Economics of Innovation and Technological Change. Blackwell.",
    "Hall, B. H., Jaffe, A., & Trajtenberg, M. (2005). Market value and patent citations. RAND Journal of Economics, 36(1), 16–38.",
    "Hall, B. H., Mairesse, J., & Mohnen, P. (2010). Measuring the returns to R&D. In Handbook of the Economics of Innovation, Vol. 2. Elsevier.",
    "Hanushek, E. A., & Woessmann, L. (2008). The role of cognitive skills in economic development. Journal of Economic Literature, 46(3), 607–668.",
    "Information Technology and Innovation Foundation (ITIF). (2025). The Bayh-Dole Act’s role in stimulating university-led regional economic growth.",
    "Jones, C. I. (1995). R&D-based models of economic growth. Journal of Political Economy, 103(4), 759–784.",
    "Jones, C. I. (2019). Paul Romer: Ideas, nonrivalry, and endogenous growth. Scandinavian Journal of Economics, 121(3), 859–883.",
    "Jones, C. I., & Williams, J. C. (1998). Measuring the social return to R&D. Quarterly Journal of Economics, 113(4), 1119–1135.",
    "Kline, S. J., & Rosenberg, N. (1986). An overview of innovation. In R. Landau & N. Rosenberg (Eds.), The positive sum strategy. National Academies Press.",
    "KISTEP. (2024). 2024년 국가연구개발사업 평가 및 조정 결과 보고서. 한국과학기술기획평가원.",
    "KISTEP. (2025). 2025 INNOPOLIS Global Forum 자료집. 한국과학기술기획평가원.",
    "Lev, B., & Gu, F. (2016). The end of accounting and the path forward for investors and managers. Wiley.",
    "Lundvall, B.-Å. (1992). National systems of innovation: Toward a theory of innovation and interactive learning. Pinter.",
    "Mankiw, N. G., Romer, D., & Weil, D. N. (1992). A contribution to the empirics of economic growth. Quarterly Journal of Economics, 107(2), 407–437.",
    "Mansfield, E. (1977). Social and private rates of return from industrial innovations. Quarterly Journal of Economics, 91(2), 221–240.",
    "Mansfield, E. (1991). Academic research and industrial innovation. Research Policy, 20(1), 1–12.",
    "Mansfield, E. (1998). Academic research and industrial innovation: An update of empirical findings. Research Policy, 26(7-8), 773–776.",
    "Mazzucato, M. (2013). The entrepreneurial state: Debunking public vs. private sector myths. Anthem Press.",
    "Mazzucato, M. (2021). Mission economy: A moonshot guide to changing capitalism. Allen Lane.",
    "Nelson, R. R. (1993). National innovation systems: A comparative analysis. Oxford University Press.",
    "Nelson, R. R., & Winter, S. G. (1982). An evolutionary theory of economic change. Belknap/Harvard University Press.",
    "Nonaka, I., & Takeuchi, H. (1995). The knowledge-creating company. Oxford University Press.",
    "OECD. (2015). Frascati Manual 2015: Guidelines for collecting and reporting data on research and experimental development. OECD Publishing.",
    "OECD. (2018). Oslo Manual 2018: Guidelines for collecting, reporting and using data on innovation, 4th edition. OECD Publishing.",
    "OECD. (2025). Challenges and opportunities of mission-oriented innovation policy in Korea. OECD Publishing.",
    "Phelps, E. S. (1966). Models of technical progress and the golden rule of research. Review of Economic Studies, 33(2), 133–145.",
    "Polanyi, M. (1967). The tacit dimension. Routledge & Kegan Paul.",
    "Rochet, J.-C., & Tirole, J. (2003). Platform competition in two-sided markets. Journal of the European Economic Association, 1(4), 990–1029.",
    "Romer, P. M. (1986). Increasing returns and long-run growth. Journal of Political Economy, 94(5), 1002–1037.",
    "Romer, P. M. (1990). Endogenous technological change. Journal of Political Economy, 98(5, Part 2), S71–S102.",
    "Scherer, F. M., & Harhoff, D. (2000). Technology policy for a world of skew-distributed outcomes. Research Policy, 29(4-5), 559–566.",
    "Schumpeter, J. A. (1912). Theorie der wirtschaftlichen Entwicklung. Duncker & Humblot.",
    "Schumpeter, J. A. (1942). Capitalism, socialism and democracy. Harper & Brothers.",
    "Smith, A. (1776). An inquiry into the nature and causes of the wealth of nations. Strahan and Cadell.",
    "Solow, R. M. (1956). A contribution to the theory of economic growth. Quarterly Journal of Economics, 70(1), 65–94.",
    "Solow, R. M. (1957). Technical change and the aggregate production function. Review of Economics and Statistics, 39(3), 312–320.",
    "Sveikauskas, L. (2007). R&D and productivity growth: A review of the literature. BLS Working Papers, 408.",
    "Trajtenberg, M. (1990). A penny for your quotes: Patent citations and the value of innovations. RAND Journal of Economics, 21(1), 172–187.",
    "Uzawa, H. (1965). Optimal technical change in an aggregative model of economic growth. International Economic Review, 6(1), 18–31.",
    "Weitzman, M. L. (1998). Recombinant growth. Quarterly Journal of Economics, 113(2), 331–360.",
    "World Bank. (2024). World Development Indicators 2024. World Bank.",
    "Zahra, S. A., & George, G. (2002). Absorptive capacity: A review, reconceptualization, and extension. Academy of Management Review, 27(2), 185–203.",
    "과학기술정보통신부. (2024). 2024년 과학기술 통계백서. 과학기술정보통신부.",
    "과학기술정책연구원(STEPI). (2024). 한국 혁신정책 50년: 회고와 전망. STEPI.",
    "교육부·한국연구재단. (2024). 2024 BK21 사업 평가 보고서. 한국연구재단.",
    "산업통상자원부. (2025). 12대 국가전략기술 R&D 추진계획. 산업통상자원부.",
    "유엔교육과학문화기구(UNESCO). (2025). Frequently asked questions (FAQs) on R&D statistics 2025. UNESCO Institute for Statistics.",
]

for ref in references:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    pf.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_run_font(run, size_pt=BODY_PT)

# ===================== 부록 A =====================
add_heading(doc, "부록 A. 주요 개념·약어 정리", level=1, page_break=True)

abbrev = [
    ("AUTM", "Association of University Technology Managers (미국)"),
    ("Bayh-Dole Act", "특허·상표법 개정법(1980, 미국)"),
    ("BERD", "Business Enterprise R&D, 기업 R&D 지출"),
    ("CDMA", "Code Division Multiple Access, 코드분할다중접속"),
    ("CHIPS Act", "미국 「반도체 및 과학법」(2022)"),
    ("DARPA", "미 국방고등연구계획국"),
    ("ETRI", "한국전자통신연구원"),
    ("GERD", "Gross Domestic Expenditure on R&D, 국내 총연구개발비"),
    ("GPT(범용)", "General-Purpose Technology, 일반목적기술"),
    ("Horizon Europe", "EU의 2021–2027 연구혁신 프레임프로그램"),
    ("IRA", "미국 「인플레이션 감축법」(2022)"),
    ("KISTEP", "한국과학기술기획평가원"),
    ("KIST", "한국과학기술연구원"),
    ("LLM", "Large Language Model, 대규모 언어모델"),
    ("MOIP", "Mission-Oriented Innovation Policy, 임무지향형 혁신정책"),
    ("NIS", "National Innovation System, 국가혁신체제"),
    ("NIH", "미 국립보건원"),
    ("NSF", "미 국립과학재단"),
    ("NTIS", "국가과학기술지식정보서비스"),
    ("OECD MSTI", "OECD Main Science and Technology Indicators"),
    ("RRI", "Responsible Research and Innovation, 책임 있는 연구와 혁신"),
    ("SNA", "System of National Accounts, 국민계정체계"),
    ("STEPI", "과학기술정책연구원"),
    ("TDX", "Time Division Exchange, 전전자교환기"),
    ("TLO", "Technology Licensing Office"),
    ("TTO", "Technology Transfer Office, 대학기술이전조직"),
    ("USPTO", "미 특허상표청"),
    ("흡수역량", "Absorptive Capacity (Cohen & Levinthal, 1990)"),
    ("창조적 파괴", "Creative Destruction (Schumpeter, 1942)"),
    ("내생적 성장이론", "Endogenous Growth Theory (Romer, 1986, 1990)"),
]

for term, desc in abbrev:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(2)
    run = p.add_run(f"  • {term}  :  ")
    set_run_font(run, size_pt=BODY_PT, bold=True)
    run = p.add_run(desc)
    set_run_font(run, size_pt=BODY_PT)

# ===================== 부록 B =====================
add_heading(doc, "부록 B. 추가 통계 표", level=1, page_break=True)

add_table_caption(doc, "B.1", "주요국 R&D 인력 1인당 R&D 지출(2022, 단위: 천 USD PPP)")
add_table(
    doc,
    header=["국가", "1인당 R&D 지출", "R&D 인력(천 명)"],
    rows=[
        ["미국", "약 372", "약 1,640"],
        ["대한민국", "약 256", "약 542"],
        ["일본", "약 268", "약 695"],
        ["독일", "약 312", "약 459"],
        ["중국", "약 117", "약 6,353"],
        ["이스라엘", "약 296", "약 117"],
        ["OECD 평균", "약 263", "—"],
    ],
    col_widths_cm=[4.0, 4.5, 4.5],
)

doc.add_paragraph()

add_table_caption(doc, "B.2", "한국 국가전략기술 12대 분야(2022 지정)")
add_table(
    doc,
    header=["No.", "분야", "비고"],
    rows=[
        ["1", "반도체·디스플레이", "메모리 1위 유지, 시스템반도체 추격"],
        ["2", "이차전지", "삼원계·LFP·차세대 셀"],
        ["3", "첨단 모빌리티", "자율주행·UAM·전기차"],
        ["4", "차세대 원자력", "SMR·핵융합"],
        ["5", "첨단 바이오", "유전자·세포치료제·신약"],
        ["6", "우주항공·해양", "발사체·위성·차세대 모빌리티"],
        ["7", "수소", "생산·저장·활용"],
        ["8", "사이버보안", "AI 활용 사이버 방어"],
        ["9", "인공지능", "파운데이션 모델·신뢰성"],
        ["10", "차세대 통신", "5G·6G·위성통신"],
        ["11", "첨단 로봇·제조", "지능형 로봇·디지털 트윈"],
        ["12", "양자", "양자컴퓨팅·통신·센싱"],
    ],
    col_widths_cm=[1.2, 4.5, 7.3],
)

# ===================== 감사의 글 =====================
add_heading(doc, "감 사 의 글", level=1, page_break=True)
ack = [
    "본 논문이 완성되기까지 여러 분의 도움이 있었습니다. 학위과정 내내 학문적 엄밀성과 사회적 책임을 함께 강조해 주신 지도교수님께 깊이 감사드립니다. 매주의 세미나에서 거친 사유를 다듬어 주신 동료 박사과정 선후배들과, 본 논문의 사례 부분에 살아 있는 데이터를 제공해 준 산업 현장의 실무자분들께도 마음 깊이 감사드립니다.",
    "또한 광주과학기술원 도서관, 한국과학기술기획평가원 자료실, OECD STI Statistics Portal, UNESCO UIS Database, 그리고 한국 NTIS 시스템의 공개 데이터는 본 연구의 통계적 토대가 되었습니다. 자료를 정리·공개해 온 모든 분들께 감사드립니다.",
    "마지막으로, 학위과정의 긴 시간을 변함없이 응원해 준 가족에게 가장 깊은 감사의 마음을 전합니다. 본 논문의 모든 미흡함은 전적으로 저자의 책임입니다.",
]
for p in ack:
    add_paragraph(doc, p, indent_first=True)

# ===================== 약력 =====================
add_heading(doc, "약       력", level=1, page_break=True)

bio = [
    ("성       명", "[저자명 기재]"),
    ("생  년  월  일", "[기재]"),
    ("학        력",
     "[학사 / 석사 / 박사 과정 기재]\n광주과학기술원 기술경영학과 박사과정"),
    ("연  구  관  심",
     "기술혁신 정책, 국가혁신체제, 임무지향 R&D, 책임 있는 혁신, "
     "디지털 전환과 거시경제"),
    ("주 요  논 문",
     "[학회·학술지 기재]"),
    ("학  술  활  동",
     "[학회 회원 / 위원 활동 기재]"),
    ("이  메  일", "[연락처 기재]"),
]
for label, val in bio:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.6
    pf.space_after = Pt(4)
    run = p.add_run(f"  ◦  {label}\t: ")
    set_run_font(run, size_pt=BODY_PT, bold=True)
    run = p.add_run(val)
    set_run_font(run, size_pt=BODY_PT)
    pf.tab_stops.add_tab_stop(Cm(4.5))

# ───────────────────────── 저장 ─────────────────────────
import os
import sys

OUT_PATH = sys.argv[1] if len(sys.argv) > 1 else "thesis.docx"
doc.save(OUT_PATH)
print(f"saved: {OUT_PATH} ({os.path.getsize(OUT_PATH):,} bytes)")
